"""Process notes from documentation sheets, cell comments and sibling documents."""

from __future__ import annotations

import zipfile
from pathlib import Path
from typing import Any

import pytest
from docx import Document as NewDocument
from openpyxl import Workbook
from openpyxl.comments import Comment

from kedge.analysis.docs import (
    _MAX_SIDECAR_NOTES,
    _MAX_TEXT_BYTES,
    MAX_NOTE_CHARS,
    _decode_text,
    _drop_partial_tail,
    extract_notes,
    sidecar_documents,
)
from kedge.analysis.model import FindingKind, Severity
from kedge.analysis.workbook import open_workbook

FIXTURES = Path(__file__).resolve().parents[1] / "fixtures"

_PROSE = (
    "Each month the exposures hand-in arrives from the credit risk team as a CSV "
    "and is pasted into the Data tab without modification."
)
_SECOND_PROSE = (
    "Collateral haircuts are applied from the reference table. Any counterparty missing "
    "a haircut is escalated to the desk before the run is signed off."
)


def _workbook(tmp_path: Path, *, name: str = "process.xlsx") -> Path:
    """A workbook with a documentation tab, a calculation tab, and one cell comment."""
    book = Workbook()

    notes = book.active
    notes.title = "Process Notes"
    notes["A1"] = "Monthly RWA process"
    notes["A2"] = _PROSE
    notes["A4"] = "Haircuts"
    notes["A5"] = _SECOND_PROSE

    data = book.create_sheet("Calc")
    data["A1"] = "ead"
    data["A2"] = 100
    data["B2"] = "=A2*2"
    data["B2"].comment = Comment("Doubled at the desk's request, agreed 2024-03.", "A. Analyst")

    path = tmp_path / name
    book.save(path)
    return path


def _write_docx(path: Path) -> None:
    document = NewDocument()
    document.add_heading("Purpose", level=1)
    document.add_paragraph("Restate the monthly capital position for the trading book.")
    document.add_heading("Steps", level=2)
    document.add_paragraph("Load the hand-in.")
    document.add_paragraph("Apply the haircut table.")
    document.save(path)


# ── documentation sheets ─────────────────────────────────────────────────────────────────────


def test_prose_is_stitched_into_blocks_rather_than_one_note_per_cell(tmp_path: Path) -> None:
    with open_workbook(_workbook(tmp_path)) as handle:
        notes, findings = extract_notes(handle)

    sheet_notes = [note for note in notes if note.source == "sheet"]
    assert findings == []
    assert len(sheet_notes) == 2, "two blank-row-separated blocks, not five cells"
    assert sheet_notes[0].origin == "Process Notes"
    assert sheet_notes[0].heading == "Monthly RWA process"
    assert sheet_notes[0].text == _PROSE
    assert sheet_notes[1].heading == "Haircuts"
    assert sheet_notes[1].location is not None


def test_calculation_sheets_contribute_no_prose(tmp_path: Path) -> None:
    with open_workbook(_workbook(tmp_path)) as handle:
        notes, _ = extract_notes(handle)

    assert not [note for note in notes if note.source == "sheet" and note.origin == "Calc"]


# ── cell comments ────────────────────────────────────────────────────────────────────────────


def test_cell_comments_become_located_notes(tmp_path: Path) -> None:
    with open_workbook(_workbook(tmp_path)) as handle:
        notes, _ = extract_notes(handle)

    comments = [note for note in notes if note.source == "cell_comment"]
    assert len(comments) == 1
    assert comments[0].origin == "Calc"
    assert comments[0].location == "B2"
    assert "agreed 2024-03" in comments[0].text
    assert comments[0].heading == "A. Analyst"


def test_a_workbook_without_comments_costs_nothing_and_says_nothing(tmp_path: Path) -> None:
    book = Workbook()
    book.active["A1"] = 1
    path = tmp_path / "plain.xlsx"
    book.save(path)

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    assert [note for note in notes if note.source == "cell_comment"] == []
    assert findings == []


# ── sibling Word documents ───────────────────────────────────────────────────────────────────


def test_a_sibling_docx_contributes_notes_with_its_own_headings(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    _write_docx(tmp_path / "process - procedure.docx")

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    docx_notes = [note for note in notes if note.source == "docx"]
    assert findings == []
    assert [note.heading for note in docx_notes] == ["Purpose", "Steps"]
    assert "Restate the monthly capital position" in docx_notes[0].text
    assert "Apply the haircut table." in docx_notes[1].text
    assert docx_notes[0].origin.endswith("process - procedure.docx")


def test_a_legacy_doc_fails_loudly_with_a_conversion_hint(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    legacy = tmp_path / "process - procedure.doc"
    legacy.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00\xff" * 64)

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    stubs = [note for note in notes if note.source == "doc_stub"]
    assert len(stubs) == 1, "a .doc must never be skipped silently"
    assert "convert" in stubs[0].text.casefold()

    assert len(findings) == 1
    assert findings[0].kind is FindingKind.UNSUPPORTED_FORMAT
    assert findings[0].severity is Severity.WARNING
    assert "procedure.doc" in (findings[0].location or "")
    assert ".docx" in (findings[0].remediation or "")


def test_a_doc_that_is_really_a_docx_is_read_anyway(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    renamed = tmp_path / "process - procedure.doc"
    _write_docx(renamed)

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    assert [note for note in notes if note.source == "doc_stub"] == []
    assert [note for note in notes if note.source == "docx"]
    assert findings == []


def test_a_corrupt_docx_produces_a_finding_rather_than_a_crash(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    (tmp_path / "process - procedure.docx").write_bytes(b"PK\x03\x04 and then nothing valid")

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    assert [note for note in notes if note.source == "docx"] == []
    assert len(findings) == 1
    assert findings[0].kind is FindingKind.UNPARSEABLE_PART
    assert "process - procedure.docx" in findings[0].message


def test_an_empty_docx_is_read_without_notes_or_findings(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    NewDocument().save(tmp_path / "process - procedure.docx")

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    assert [note for note in notes if note.source == "docx"] == []
    assert findings == []


# ── sibling markdown ─────────────────────────────────────────────────────────────────────────

_MARKDOWN = """\
# Monthly RWA process

Restate the monthly capital position for the trading book.

## Steps

Load the hand-in.
Apply the haircut table.

Notes on exceptions
-------------------

Any counterparty missing a haircut is escalated to the desk.
"""


def test_a_sibling_markdown_file_is_split_on_its_own_headings(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    (tmp_path / "process - procedure.md").write_text(_MARKDOWN, encoding="utf-8")

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    markdown = [note for note in notes if note.source == "markdown"]
    assert findings == []
    assert [note.heading for note in markdown] == [
        "Monthly RWA process",
        "Steps",
        "Notes on exceptions",
    ], "ATX and setext headings both carry the author's structure"
    assert "Restate the monthly capital position" in markdown[0].text
    assert "Apply the haircut table." in markdown[1].text
    assert markdown[0].origin.endswith("process - procedure.md")


def test_markdown_notes_are_located_by_line_number(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    (tmp_path / "process - procedure.md").write_text(_MARKDOWN, encoding="utf-8")

    with open_workbook(path) as handle:
        notes, _ = extract_notes(handle)

    markdown = [note for note in notes if note.source == "markdown"]
    assert markdown[0].location == "line 3", "one body line under the first heading"
    assert markdown[1].location == "lines 7-8"


def test_a_heading_inside_a_fenced_code_block_does_not_split_a_note(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    (tmp_path / "process - procedure.md").write_text(
        "# Rebuild\n\nRun it like this:\n\n```bash\n# refresh the cache first\nmake refresh\n```\n",
        encoding="utf-8",
    )

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    markdown = [note for note in notes if note.source == "markdown"]
    assert findings == []
    assert [note.heading for note in markdown] == ["Rebuild"], "a shell comment is not a heading"
    assert "refresh the cache first" in markdown[0].text


# ── sibling plain text ───────────────────────────────────────────────────────────────────────


def test_a_sibling_text_file_is_split_on_blank_lines_without_inventing_headings(
    tmp_path: Path,
) -> None:
    path = _workbook(tmp_path)
    (tmp_path / "process - procedure.txt").write_text(
        "Check the FX rate.\nIt comes from the rates tab.\n\nRun the model.\n",
        encoding="utf-8",
    )

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    text_notes = [note for note in notes if note.source == "plain_text"]
    assert findings == []
    assert len(text_notes) == 2, "two blank-line-separated paragraphs"
    assert text_notes[0].text == "Check the FX rate.\nIt comes from the rates tab."
    assert text_notes[0].location == "lines 1-2"
    assert all(note.heading is None for note in text_notes), "a .txt has no heading vocabulary"


def test_a_short_paragraph_in_a_text_file_is_kept(tmp_path: Path) -> None:
    """The sheet reader drops a stray label. Nothing in a text file is stray."""
    path = _workbook(tmp_path)
    (tmp_path / "process - procedure.txt").write_text(
        "Check the rate.\n\nRun it.\n\nSign off.\n", encoding="utf-8"
    )

    with open_workbook(path) as handle:
        notes, _ = extract_notes(handle)

    assert len([note for note in notes if note.source == "plain_text"]) == 3


def test_a_text_sidecar_that_will_not_decode_produces_a_finding_rather_than_a_traceback(
    tmp_path: Path,
) -> None:
    """`UnicodeDecodeError` is a `ValueError`, so `except OSError` would not have caught it."""
    path = _workbook(tmp_path)
    bad = tmp_path / "process - procedure.txt"
    bad.write_bytes(b"Step one\x00\x81\x8d\x8f then nothing decodable\x90\x9d")

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    assert [note for note in notes if note.source == "plain_text"] == []
    assert len(findings) == 1
    assert findings[0].kind is FindingKind.UNPARSEABLE_PART
    assert findings[0].severity is Severity.WARNING
    assert "process - procedure.txt" in findings[0].message
    assert "UTF-8" in (findings[0].remediation or "")


def test_a_text_sidecar_in_the_windows_ansi_codepage_is_still_read(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    (tmp_path / "process - procedure.txt").write_bytes(
        "Escalate to Renée before sign-off.".encode("cp1252")
    )

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    text_notes = [note for note in notes if note.source == "plain_text"]
    assert findings == []
    assert "Renée" in text_notes[0].text


def test_a_utf16_text_sidecar_is_read_from_its_byte_order_mark(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    (tmp_path / "process - procedure.txt").write_bytes(
        "Reconcile against last month.".encode("utf-16")
    )

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    text_notes = [note for note in notes if note.source == "plain_text"]
    assert findings == []
    assert text_notes[0].text == "Reconcile against last month."


# ── caps ─────────────────────────────────────────────────────────────────────────────────────


def test_an_oversized_text_sidecar_is_read_only_to_the_byte_cap(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    body = ("The desk reviews every exception before the run is signed off.\n\n" * 40_000)[
        : _MAX_TEXT_BYTES * 2
    ]
    sidecar = tmp_path / "process - procedure.txt"
    sidecar.write_text(body, encoding="utf-8")
    assert sidecar.stat().st_size > _MAX_TEXT_BYTES, "the fixture has to exceed the cap"

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    text_notes = [note for note in notes if note.source == "plain_text"]
    assert findings == []
    assert len(text_notes) <= _MAX_SIDECAR_NOTES
    assert sum(len(note.text) for note in text_notes) < _MAX_TEXT_BYTES


def test_no_more_than_the_note_cap_survives_one_sidecar(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    sections = "".join(f"## Step {n}\n\nDo the thing.\n\n" for n in range(_MAX_SIDECAR_NOTES + 50))
    (tmp_path / "process - procedure.md").write_text(sections, encoding="utf-8")

    with open_workbook(path) as handle:
        notes, _ = extract_notes(handle)

    assert len([note for note in notes if note.source == "markdown"]) == _MAX_SIDECAR_NOTES


def test_a_long_text_note_is_truncated_at_the_per_note_ceiling(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    (tmp_path / "process - procedure.txt").write_text("word " * 2_000, encoding="utf-8")

    with open_workbook(path) as handle:
        notes, _ = extract_notes(handle)

    text_notes = [note for note in notes if note.source == "plain_text"]
    assert len(text_notes[0].text) <= MAX_NOTE_CHARS
    assert text_notes[0].text.endswith("truncated]")


def test_the_byte_cap_never_leaves_half_a_character_behind() -> None:
    """The cap cuts at a byte offset, so the tail it leaves has to be made decodable."""
    for tail in ("é", "€", "\U0001f600", "x"):  # two-, three- and four-byte sequences
        whole = f"escalate {tail}".encode()
        for cut in range(1, len(whole) + 1):
            _drop_partial_tail(whole[:cut]).decode()  # strict: raises if still broken


def test_a_file_that_cp1252_turns_into_nulls_is_reported_not_believed() -> None:
    """UTF-16 with no mark decodes 'fine' as cp1252, into mojibake. That is not text."""
    assert _decode_text("Reconcile.".encode("utf-16-le")) is None
    assert _decode_text(b"Reconcile.") == "Reconcile."


def test_no_more_than_ten_sidecars_are_attached_to_one_workbook(tmp_path: Path) -> None:
    workbook = tmp_path / "rwa_monthly.xlsx"
    workbook.touch()
    for index in range(20):
        (tmp_path / f"rwa_monthly - procedure {index:02d}.md").touch()

    assert len(sidecar_documents(workbook)) == 10, "the cap is shared across every format"


# ── sidecar selection ────────────────────────────────────────────────────────────────────────


def test_sidecar_selection_matches_by_stem_or_convention(tmp_path: Path) -> None:
    workbook = tmp_path / "rwa_monthly.xlsx"
    workbook.touch()
    for name in (
        "rwa_monthly.docx",
        "rwa_monthly - procedure.docx",
        "Process Notes.docx",
        "somebody elses report.docx",
        "~$rwa_monthly.docx",
    ):
        (tmp_path / name).touch()

    found = {path.name for path in sidecar_documents(workbook)}

    assert found == {"rwa_monthly.docx", "rwa_monthly - procedure.docx", "Process Notes.docx"}
    assert "somebody elses report.docx" not in found, "unrelated documents must not be attached"
    assert "~$rwa_monthly.docx" not in found, "Word lock files are not documents"


def test_sidecar_selection_accepts_markdown_and_plain_text(tmp_path: Path) -> None:
    workbook = tmp_path / "rwa_monthly.xlsx"
    workbook.touch()
    for name in (
        "rwa_monthly.md",
        "monthly-process.md",
        "runbook.txt",
        "exposures export.txt",
        "rwa_monthly.csv",
    ):
        (tmp_path / name).touch()

    found = {path.name for path in sidecar_documents(workbook)}

    assert found == {"rwa_monthly.md", "monthly-process.md", "runbook.txt"}
    assert "exposures export.txt" not in found, "an unrelated .txt is not a procedure"
    assert "rwa_monthly.csv" not in found, "data beside the workbook is not documentation"


# ── attachment by guesswork ──────────────────────────────────────────────────────────────────


def test_a_document_attached_by_its_filename_alone_says_so(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    _write_docx(tmp_path / "SOP.docx")

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    assert [note for note in notes if note.source == "docx"], "the guess is still made"
    assert len(findings) == 1
    assert findings[0].kind is FindingKind.DOCUMENT_ATTACHED_BY_FILENAME
    assert findings[0].severity is Severity.INFO, "a guess worth stating is not an error"
    assert "SOP.docx" in findings[0].message
    assert "process.xlsx" in findings[0].message
    assert "process - SOP.docx" in (findings[0].remediation or ""), "name the exact rename"


def test_a_document_sharing_the_workbook_stem_is_attached_silently(tmp_path: Path) -> None:
    """A stem match is a statement about this workbook, not a guess. No finding."""
    path = _workbook(tmp_path)
    _write_docx(tmp_path / "process - procedure.docx")

    with open_workbook(path) as handle:
        _, findings = extract_notes(handle)

    assert findings == []


def test_the_guess_is_recorded_for_markdown_too(tmp_path: Path) -> None:
    path = _workbook(tmp_path)
    (tmp_path / "runbook.md").write_text("# Runbook\n\nRestate the position.\n", encoding="utf-8")

    with open_workbook(path) as handle:
        notes, findings = extract_notes(handle)

    assert [note for note in notes if note.source == "markdown"]
    assert [finding.kind for finding in findings] == [FindingKind.DOCUMENT_ATTACHED_BY_FILENAME]


def test_a_guessed_attachment_that_is_also_unreadable_reports_both(tmp_path: Path) -> None:
    """The guess is recorded whether or not the document then reads."""
    path = _workbook(tmp_path)
    (tmp_path / "SOP.doc").write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00\xff" * 64)

    with open_workbook(path) as handle:
        _, findings = extract_notes(handle)

    assert [finding.kind for finding in findings] == [
        FindingKind.DOCUMENT_ATTACHED_BY_FILENAME,
        FindingKind.UNSUPPORTED_FORMAT,
    ]


def test_sidecar_selection_on_a_missing_directory_is_empty() -> None:
    assert sidecar_documents(Path("does") / "not" / "exist" / "book.xlsx") == []


# ── degradation ──────────────────────────────────────────────────────────────────────────────


class _BrokenHandle:
    """A handle that fails at every turn. Every extractor must survive it."""

    def __init__(self, path: Path) -> None:
        self.path = path

    def __getattr__(self, name: str) -> Any:
        message = f"deliberately broken handle: {name}"
        raise RuntimeError(message)


def test_a_handle_that_raises_on_everything_yields_empty_results(tmp_path: Path) -> None:
    # An empty directory, not a relative path: sidecars are found beside the workbook, so a
    # handle pointing at `nowhere.xlsx` scans the working directory and picks up whatever
    # happens to be in it -- the repo's own README.md, when the suite runs from the root.
    notes, findings = extract_notes(_BrokenHandle(tmp_path / "nowhere.xlsx"))

    assert notes == []
    assert findings == []


def test_a_handle_with_no_views_yields_empty_results(tmp_path: Path) -> None:
    class _Bare:
        path = tmp_path / "nowhere.xlsx"

    notes, findings = extract_notes(_Bare())

    assert (notes, findings) == ([], [])


# ── corpus fixture ───────────────────────────────────────────────────────────────────────────


@pytest.mark.corpus
def test_documented_fixture_yields_prose_a_comment_and_the_word_procedure() -> None:
    fixture = FIXTURES / "documented.xlsx"
    if not fixture.is_file():
        pytest.skip("tests/fixtures/documented.xlsx has not landed yet")

    with open_workbook(fixture) as handle:
        notes, findings = extract_notes(handle)

    sources = {note.source for note in notes}
    assert "sheet" in sources, "the documented fixture is meant to carry a notes tab"
    if (FIXTURES / "documented_procedure.docx").is_file():
        assert "docx" in sources
    if (FIXTURES / "procedure_legacy.doc").is_file():
        assert "doc_stub" in sources
        assert any(finding.kind is FindingKind.UNSUPPORTED_FORMAT for finding in findings)


@pytest.mark.corpus
def test_the_documented_fixture_warns_about_the_guess_and_not_about_the_stem_match() -> None:
    """The whole point of the finding, on the one fixture that shows both sides of it.

    `documented_procedure.docx` starts with the workbook's own stem, so it was attached on
    evidence about *this* workbook and says nothing. `procedure_legacy.doc` shares nothing
    with it and was attached on the naming convention alone, so it does.
    """
    fixture = FIXTURES / "documented.xlsx"
    if not fixture.is_file():
        pytest.skip("tests/fixtures/documented.xlsx has not landed yet")

    with open_workbook(fixture) as handle:
        _, findings = extract_notes(handle)

    guessed = {
        Path(finding.location or "").name
        for finding in findings
        if finding.kind is FindingKind.DOCUMENT_ATTACHED_BY_FILENAME
    }

    assert "procedure_legacy.doc" in guessed, "a name match is a guess and must be recorded"
    assert "documented_procedure.docx" not in guessed, (
        "a stem match is evidence about this workbook, not a guess"
    )


@pytest.mark.corpus
def test_the_hostile_fixture_never_raises() -> None:
    fixture = FIXTURES / "hostile.xlsx"
    if not fixture.is_file():
        pytest.skip("tests/fixtures/hostile.xlsx has not landed yet")

    with zipfile.ZipFile(fixture):
        pass  # it is at least a readable archive
    with open_workbook(fixture) as handle:
        notes, findings = extract_notes(handle)

    assert isinstance(notes, list)
    assert isinstance(findings, list)
