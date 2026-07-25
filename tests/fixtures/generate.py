#!/usr/bin/env python
"""Generate the committed .xlsx fixture corpus for kedge's M1 analyser.

Run from the repository root::

    uv run python tests/fixtures/generate.py
    uv run --with pywin32 python tests/fixtures/generate.py --verify-with-excel

This script is the source of truth for every file in ``tests/fixtures/``. The
workbooks are committed so that the test suite has no generation step, but they
must always be reproducible from here — ``tests/unit/test_fixtures_generate.py``
regenerates the corpus into a temporary directory and asserts the committed
files match part for part.

Determinism is deliberate and load-bearing. All randomness is seeded, document
properties are pinned to a fixed timestamp, and every archive is rewritten
through :func:`write_parts` with a fixed zip entry date so that regenerating
produces byte-identical output. See the "Determinism" section of README.md.

Three of the fixtures need XML that openpyxl will not write: the Power Query
``DataMashup`` part, ``xl/connections.xml``, and the cached calculated values in
``clean_pipeline.xlsx``. Those are built by post-processing the saved archive and
are each verified by reading them back — see the ``verify_*`` functions.
"""

from __future__ import annotations

import argparse
import base64
import datetime as dt
import io
import logging
import random
import re
import struct
import sys
import zipfile
from dataclasses import dataclass
from decimal import ROUND_HALF_UP, Decimal
from pathlib import Path
from typing import Any

from openpyxl import Workbook, load_workbook
from openpyxl.comments import Comment
from openpyxl.packaging.relationship import Relationship
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.workbook.defined_name import DefinedName
from openpyxl.workbook.external_link.external import (
    ExternalBook,
    ExternalLink,
    ExternalSheetNames,
)
from openpyxl.worksheet.worksheet import Worksheet

FIXTURE_DIR = Path(__file__).parent
SEED = 20260724

# Pinned so that docProps/core.xml carries no wall-clock time.
FIXED_TIMESTAMP = dt.datetime(2026, 1, 1, 0, 0, 0)
# Pinned zip entry date; 1980-01-01 is the earliest a zip can represent.
FIXED_ZIP_DATE = (1980, 1, 1, 0, 0, 0)

SHEET_MAIN_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
OFFICE_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
DATAMASHUP_NS = "http://schemas.microsoft.com/DataMashup"

logger = logging.getLogger("kedge.fixtures")


# =============================================================================
# Excel value semantics
# =============================================================================
# The cached values written into clean_pipeline.xlsx must be the numbers Excel
# would have computed, not the numbers Python would. The only divergence that
# bites at this scale is the rounding mode (PLAN section 2.6): Excel rounds half
# away from zero, Python and polars round half to even.
#
# This is deliberately a local three-line helper rather than an import from
# kedge.xl. The fixtures are the corpus that kedge.xl is validated against, so
# they must not depend on it; if both were wrong in the same way the test would
# still pass. The Excel COM verification path (--verify-with-excel) is the
# independent oracle that keeps this honest.


def excel_round(value: float, digits: int) -> float:
    """Round half away from zero, the way Excel's ROUND does.

    Two behaviours have to line up, and only the first is widely known:

    1. Ties go away from zero (0.5 to 1, 2.5 to 3), not to even.
    2. Excel first collapses the operand to 15 significant decimal digits, then
       rounds. A product that lands at 4491760.574999999... in binary double
       precision is 4491760.575 to Excel, which then rounds up to 4491760.58.
       Rounding the raw double instead gives 4491760.57.

    Point 2 was found by comparing this corpus against a real Excel
    recalculation (``--verify-with-excel``): it moved 2 of 2500 cached values,
    which then propagated into 2 of the 6 Output totals. It is worth knowing
    about in ``kedge.xl`` for the same reason.
    """
    significant = Decimal(f"{value:.15g}")
    quantum = Decimal(1).scaleb(-digits)
    return float(significant.quantize(quantum, rounding=ROUND_HALF_UP))


# =============================================================================
# Archive plumbing
# =============================================================================


def read_parts(path: Path) -> dict[str, bytes]:
    """Read every part of an OPC archive into a name to bytes mapping."""
    with zipfile.ZipFile(path) as archive:
        return {name: archive.read(name) for name in archive.namelist()}


_TIMESTAMP_ELEMENT = re.compile(
    rb"(<dcterms:(?:created|modified)\b[^>]*>)[^<]*(</dcterms:(?:created|modified)>)"
)


def pin_document_timestamps(parts: dict[str, bytes]) -> None:
    """Overwrite the created and modified timestamps in ``docProps/core.xml``.

    This has to happen after the save, not before: openpyxl's ``save_workbook``
    assigns ``properties.modified = datetime.now()`` unconditionally, so any
    value set on the workbook object is discarded. python-docx behaves the same
    way. Without this the corpus is only reproducible within a single wall-clock
    second, which is exactly the sort of "passes locally, fails in CI" flake the
    determinism test exists to prevent.
    """
    key = "docProps/core.xml"
    if key not in parts:
        return
    stamp = FIXED_TIMESTAMP.strftime("%Y-%m-%dT%H:%M:%SZ").encode("ascii")
    parts[key] = _TIMESTAMP_ELEMENT.sub(rb"\g<1>" + stamp + rb"\g<2>", parts[key])


def write_parts(path: Path, parts: dict[str, bytes]) -> None:
    """Write parts back out as a normalised, deterministic zip archive.

    Entries are sorted by name and stamped with a fixed date so that two runs of
    this script produce byte-identical files. openpyxl's own ``save`` uses the
    wall clock for entry timestamps, which is why every fixture is routed
    through here even when nothing needed patching.
    """
    pin_document_timestamps(parts)
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED, allowZip64=True) as archive:
        for name in sorted(parts):
            info = zipfile.ZipInfo(name, date_time=FIXED_ZIP_DATE)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            info.create_system = 0
            archive.writestr(info, parts[name])


def normalise(path: Path) -> None:
    """Rewrite an archive in place so its bytes are reproducible."""
    write_parts(path, read_parts(path))


def add_content_type_override(parts: dict[str, bytes], part_name: str, content_type: str) -> None:
    """Register an Override for a part in ``[Content_Types].xml``."""
    key = "[Content_Types].xml"
    xml = parts[key].decode("utf-8")
    override = f'<Override PartName="{part_name}" ContentType="{content_type}"/>'
    if override in xml:
        return
    parts[key] = xml.replace("</Types>", f"{override}</Types>").encode("utf-8")


def add_workbook_relationship(
    parts: dict[str, bytes], rel_id: str, rel_type: str, target: str
) -> None:
    """Append a relationship to ``xl/_rels/workbook.xml.rels``."""
    key = "xl/_rels/workbook.xml.rels"
    xml = parts[key].decode("utf-8")
    rel = f'<Relationship Type="{rel_type}" Target="{target}" Id="{rel_id}"/>'
    parts[key] = xml.replace("</Relationships>", f"{rel}</Relationships>").encode("utf-8")


def next_workbook_rel_id(parts: dict[str, bytes]) -> str:
    """Return an unused rId for ``xl/_rels/workbook.xml.rels``."""
    xml = parts["xl/_rels/workbook.xml.rels"].decode("utf-8")
    used = {int(n) for n in re.findall(r'Id="rId(\d+)"', xml)}
    return f"rId{max(used, default=0) + 1}"


def sheet_part_map(parts: dict[str, bytes]) -> dict[str, str]:
    """Map worksheet display name to its part name inside the archive.

    Resolved through workbook.xml and its rels rather than assuming that the
    Nth sheet lives in ``sheetN.xml``, which only holds while there are no
    chartsheets or dialogue sheets.
    """
    workbook = parts["xl/workbook.xml"].decode("utf-8")
    rels = parts["xl/_rels/workbook.xml.rels"].decode("utf-8")

    rel_targets: dict[str, str] = {}
    for match in re.finditer(r"<Relationship\b[^>]*/>", rels):
        tag = match.group(0)
        rid = re.search(r'Id="([^"]+)"', tag)
        target = re.search(r'Target="([^"]+)"', tag)
        if rid and target:
            rel_targets[rid.group(1)] = target.group(1)

    mapping: dict[str, str] = {}
    for match in re.finditer(r"<sheet\b[^>]*/>", workbook):
        tag = match.group(0)
        name = re.search(r'name="([^"]+)"', tag)
        rid = re.search(r'r:id="([^"]+)"', tag)
        if not (name and rid):
            continue
        target = rel_targets[rid.group(1)]
        mapping[name.group(1)] = target.lstrip("/") if target.startswith("/") else f"xl/{target}"
    return mapping


# =============================================================================
# Cached calculated values
# =============================================================================
# openpyxl emits every formula cell as <c r="D2"><f>...</f><v></v></c> — the
# value element is present but empty, which is exactly why load_workbook(
# data_only=True) yields None for a tool-written workbook (PLAN section 1.5).
# Injecting a cached value is therefore a targeted substitution: fill in the
# empty <v>, and add t="str" when the result is text.

_FORMULA_CELL = re.compile(
    rb'<c r="(?P<ref>[A-Z]+\d+)"(?P<attrs>[^>]*)><f>(?P<formula>.*?)</f><v></v></c>'
)


def _encode_cached(value: Any) -> tuple[bytes, bytes]:
    """Return the ``<v>`` text and any extra ``<c>`` attributes for a value."""
    if isinstance(value, bool):
        return (b"1" if value else b"0"), b' t="b"'
    if isinstance(value, str):
        escaped = value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        return escaped.encode("utf-8"), b' t="str"'
    if isinstance(value, int):
        return str(value).encode("ascii"), b""
    if isinstance(value, float):
        text = repr(int(value)) if value.is_integer() else repr(value)
        return text.encode("ascii"), b""
    msg = f"cannot write {type(value).__name__} as a cached value: {value!r}"
    raise TypeError(msg)


def inject_cached_values(
    parts: dict[str, bytes], cached: dict[str, dict[str, Any]]
) -> dict[str, int]:
    """Write cached calculated values alongside the formulas in the sheet XML.

    ``cached`` maps sheet name to a mapping of cell reference to value. Returns
    the number of cells populated per sheet so the caller can assert that every
    value it supplied actually landed — a silently missed substitution would
    produce a fixture that claims cached values it does not have.
    """
    part_names = sheet_part_map(parts)
    populated: dict[str, int] = {}

    for sheet_name, values in cached.items():
        part = part_names[sheet_name]
        wanted = {ref.encode("ascii"): value for ref, value in values.items()}
        hits = 0

        def replace(match: re.Match[bytes], wanted: dict[bytes, Any] = wanted) -> bytes:
            nonlocal hits
            ref = match.group("ref")
            if ref not in wanted:
                return match.group(0)
            hits += 1
            text, extra = _encode_cached(wanted[ref])
            return (
                b'<c r="'
                + ref
                + b'"'
                + match.group("attrs")
                + extra
                + b"><f>"
                + match.group("formula")
                + b"</f><v>"
                + text
                + b"</v></c>"
            )

        parts[part] = _FORMULA_CELL.sub(replace, parts[part])
        if hits != len(wanted):
            msg = (
                f"cached value injection missed cells on sheet {sheet_name!r}: "
                f"supplied {len(wanted)}, wrote {hits}"
            )
            raise RuntimeError(msg)
        populated[sheet_name] = hits

    return populated


def set_full_calc_on_load(parts: dict[str, bytes], enabled: bool) -> None:
    """Set ``fullCalcOnLoad`` so Excel trusts (or rebuilds) the cached values."""
    key = "xl/workbook.xml"
    xml = parts[key].decode("utf-8")
    flag = "1" if enabled else "0"
    xml = re.sub(r'fullCalcOnLoad="[01]"', f'fullCalcOnLoad="{flag}"', xml)
    parts[key] = xml.encode("utf-8")


# =============================================================================
# Power Query: the MS-QDEFF DataMashup envelope
# =============================================================================
# Excel does not store Section1.m as a plain zip. It stores a binary envelope
# (MS-QDEFF) whose second field is the OPC package holding the M source, then
# base64-encodes the whole envelope into a <DataMashup> element. An extractor
# that base64-decodes and hands the result straight to zipfile will fail on the
# eight-byte header, which is precisely the behaviour this fixture exists to
# catch. See README.md for the byte layout.

PERMISSIONS_XML = (
    '<?xml version="1.0" encoding="utf-8"?>'
    '<PermissionList xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"'
    ' xmlns:xsd="http://www.w3.org/2001/XMLSchema">'
    "<CanEvaluateFuturePackages>false</CanEvaluateFuturePackages>"
    "<FirewallEnabled>true</FirewallEnabled>"
    '<WorkbookGroupType xsi:nil="true" />'
    "</PermissionList>"
)

PACKAGE_CONTENT_TYPES = (
    '<?xml version="1.0" encoding="utf-8"?>'
    '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
    '<Default Extension="xml" ContentType="text/xml" />'
    '<Default Extension="m" ContentType="application/x-ms-m" />'
    '<Override PartName="/Config/Package.xml" ContentType="application/xml" />'
    "</Types>"
)

PACKAGE_CONFIG = (
    '<?xml version="1.0" encoding="utf-8"?>'
    f'<Package xmlns="{DATAMASHUP_NS}">'
    "<Version>2.126.1076.0</Version>"
    "<MinVersion>2.21.0.0</MinVersion>"
    "<Culture>en-GB</Culture>"
    "</Package>"
)


def _metadata_xml(query_names: list[str]) -> str:
    """Build a LocalPackageMetadataFile with the entry set Excel really writes.

    The full ``StableEntries`` block matters for more than realism: it is what
    makes the metadata block big. Excel records roughly 1.5KB per query here,
    so a workbook with a few dozen queries pushes the block past 64KB — which is
    the threshold at which a naive reader stops being able to find the inner
    package. See ``build_powerquery_large_metadata``.
    """
    items = [
        "<Item><ItemLocation><ItemType>AllFormulas</ItemType>"
        "<ItemPath /><ItemName /></ItemLocation><StableEntries /></Item>"
    ]
    for index, name in enumerate(query_names):
        columns = "&quot;,&quot;".join(f"{name}_col{n}" for n in range(1, 6))
        items.append(
            "<Item><ItemLocation><ItemType>Formula</ItemType>"
            f"<ItemPath>Section1/{name}</ItemPath><ItemName /></ItemLocation>"
            "<StableEntries>"
            '<Entry Type="IsPrivate" Value="l0" />'
            '<Entry Type="FillEnabled" Value="l1" />'
            '<Entry Type="FillObjectType" Value="sConnectionOnly" />'
            f'<Entry Type="FillTarget" Value="s{name}" />'
            '<Entry Type="FillToDataModelEnabled" Value="l0" />'
            '<Entry Type="FillCount" Value="l0" />'
            '<Entry Type="FillErrorCode" Value="sUnknown" />'
            '<Entry Type="FillErrorCount" Value="l0" />'
            '<Entry Type="FillLastUpdated" Value="d2026-06-30T08:14:22.7654321Z" />'
            '<Entry Type="FillStatus" Value="sComplete" />'
            f'<Entry Type="FillColumnNames" Value="s[&quot;{columns}&quot;]" />'
            '<Entry Type="FillColumnTypes" Value="sAgQCBAI=" />'
            '<Entry Type="ResultType" Value="sTable" />'
            f'<Entry Type="QueryGroupID" Value="s{{9F1B0000-0000-4E3C-9F1B-{index:012d}}}" />'
            '<Entry Type="NameUpdatedAfterFill" Value="l0" />'
            '<Entry Type="AddedToDataModel" Value="l0" />'
            '<Entry Type="RecoveryTargetSheet" Value="sStaging!" />'
            '<Entry Type="RecoveryTargetColumn" Value="l1" />'
            '<Entry Type="RecoveryTargetRow" Value="l1" />'
            '<Entry Type="BufferNextRefresh" Value="l1" />'
            "</StableEntries></Item>"
        )
    return (
        '<?xml version="1.0" encoding="utf-8"?>'
        '<LocalPackageMetadataFile xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"'
        ' xmlns:xsd="http://www.w3.org/2001/XMLSchema">'
        f"<Items>{''.join(items)}</Items>"
        "</LocalPackageMetadataFile>"
    )


def build_mashup_package(section_m: str) -> bytes:
    """Build the inner OPC package holding ``Formulas/Section1.m``."""
    buffer = io.BytesIO()
    contents = {
        "[Content_Types].xml": PACKAGE_CONTENT_TYPES.encode("utf-8"),
        "Config/Package.xml": PACKAGE_CONFIG.encode("utf-8"),
        "Formulas/Section1.m": section_m.encode("utf-8"),
    }
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for name in sorted(contents):
            info = zipfile.ZipInfo(name, date_time=FIXED_ZIP_DATE)
            info.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(info, contents[name])
    return buffer.getvalue()


def assemble_qdeff(package: bytes, query_names: list[str]) -> bytes:
    """Wrap an arbitrary package payload in the MS-QDEFF envelope.

    Split out from :func:`build_datamashup` so the deliberately broken fixtures
    can put something other than a valid OPC package in the package slot while
    keeping the envelope itself well formed.
    """
    permissions = PERMISSIONS_XML.encode("utf-8")
    metadata = struct.pack("<I", 0) + _metadata_xml(query_names).encode("utf-8")
    bindings = b""

    return b"".join(
        [
            struct.pack("<I", 0),  # envelope version, must be 0
            struct.pack("<I", len(package)),
            package,
            struct.pack("<I", len(permissions)),
            permissions,
            struct.pack("<I", len(metadata)),
            metadata,
            struct.pack("<I", len(bindings)),
            bindings,
        ]
    )


def build_datamashup(section_m: str, query_names: list[str]) -> bytes:
    """Build a complete, well-formed MS-QDEFF envelope around ``Formulas/Section1.m``.

    The composition of :func:`build_mashup_package` and :func:`assemble_qdeff`, kept as a
    named function because it is what every healthy fixture wants. The deliberately broken
    fixtures call the two halves separately so they can put something other than a valid OPC
    package in the package slot.

    Args:
        section_m: The concatenated M source for every query.
        query_names: Query names, which the envelope's metadata block records.

    Returns:
        The raw envelope bytes, ready to be base64-encoded into a ``<DataMashup>`` element.
    """
    return assemble_qdeff(build_mashup_package(section_m), query_names)


def parse_datamashup(envelope: bytes) -> dict[str, bytes]:
    """Reverse :func:`build_datamashup`; used by the verification pass."""
    offset = 0

    def take_uint32() -> int:
        nonlocal offset
        (value,) = struct.unpack_from("<I", envelope, offset)
        offset += 4
        return value

    def take_block() -> bytes:
        nonlocal offset
        length = take_uint32()
        block = envelope[offset : offset + length]
        offset += length
        return block

    version = take_uint32()
    if version != 0:
        msg = f"unexpected DataMashup envelope version {version}, expected 0"
        raise ValueError(msg)
    return {
        "package": take_block(),
        "permissions": take_block(),
        "metadata": take_block(),
        "bindings": take_block(),
    }


CUSTOM_XML_PROPS_TYPE = "application/vnd.openxmlformats-officedocument.customXmlProperties+xml"


def _item_props(item_id: str, schema_uri: str) -> bytes:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<ds:datastoreItem ds:itemID="{item_id}"'
        f' xmlns:ds="{OFFICE_REL}/customXml">'
        f'<ds:schemaRefs><ds:schemaRef ds:uri="{schema_uri}"/></ds:schemaRefs>'
        "</ds:datastoreItem>"
    ).encode()


def _item_rels(index: int) -> bytes:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<Relationships xmlns="{PKG_REL_NS}">'
        f'<Relationship Id="rId1" Type="{OFFICE_REL}/customXmlProps"'
        f' Target="itemProps{index}.xml"/>'
        "</Relationships>"
    ).encode()


def attach_custom_xml(parts: dict[str, bytes], items: list[tuple[bytes, str, str]]) -> None:
    """Attach ``customXml/item<N>.xml`` parts and wire up types and rels.

    ``items`` is an ordered list of ``(content, schema_uri, item_guid)``; the
    index in the archive follows list position, so the caller controls which
    slot the DataMashup lands in.
    """
    for index, (content, schema_uri, item_guid) in enumerate(items, start=1):
        parts[f"customXml/item{index}.xml"] = content
        parts[f"customXml/itemProps{index}.xml"] = _item_props(item_guid, schema_uri)
        parts[f"customXml/_rels/item{index}.xml.rels"] = _item_rels(index)
        add_content_type_override(parts, f"/customXml/itemProps{index}.xml", CUSTOM_XML_PROPS_TYPE)
        add_workbook_relationship(
            parts,
            next_workbook_rel_id(parts),
            f"{OFFICE_REL}/customXml",
            f"../customXml/item{index}.xml",
        )


# =============================================================================
# Legacy connections
# =============================================================================

CONNECTIONS_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.connections+xml"


def xml_attr(text: str) -> str:
    """Escape text for an XML attribute, preserving newlines as ``&#10;``.

    XML attribute-value normalisation turns a literal newline into a space, so a
    multi-line SQL statement must encode its line breaks numerically or the
    statement comes back as one long line. Excel writes ``&#10;``; so do we.
    """
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("\r\n", "\n")
        .replace("\n", "&#10;")
    )


def attach_connections(parts: dict[str, bytes], connections_xml: bytes) -> None:
    """Attach ``xl/connections.xml`` and wire up the type and relationship."""
    parts["xl/connections.xml"] = connections_xml
    add_content_type_override(parts, "/xl/connections.xml", CONNECTIONS_TYPE)
    add_workbook_relationship(
        parts, next_workbook_rel_id(parts), f"{OFFICE_REL}/connections", "connections.xml"
    )


# =============================================================================
# Shared workbook helpers
# =============================================================================

HEADER_FONT = Font(bold=True)
HEADER_FILL = PatternFill("solid", fgColor="DDDDDD")


def new_workbook() -> Workbook:
    """Create a workbook with pinned document properties."""
    workbook = Workbook()
    workbook.properties.created = FIXED_TIMESTAMP
    workbook.properties.modified = FIXED_TIMESTAMP
    workbook.properties.creator = "kedge fixture generator"
    workbook.properties.lastModifiedBy = "kedge fixture generator"
    return workbook


def write_header(sheet: Worksheet, row: int, headers: list[str]) -> None:
    """Write a bold header row."""
    for column, name in enumerate(headers, start=1):
        cell = sheet.cell(row=row, column=column, value=name)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL


def save(workbook: Workbook, path: Path) -> None:
    """Save and normalise, leaving a reproducible archive on disk."""
    workbook.save(path)
    normalise(path)


# =============================================================================
# Fixture 1: clean_pipeline.xlsx
# =============================================================================

COUNTERPARTIES = [
    "Alderney Capital",
    "Brightwater LLP",
    "Cinnabar Holdings",
    "Dunlin Asset Mgmt",
    "Everton Securities",
    "Fenchurch Partners",
    "Greenwich Trading",
    "Harlow Investments",
    "Ilkley Financial",
    "Jarrow Markets",
    "Kelso Brothers",
    "Lambourn Group",
]
ASSET_CLASSES = ["Equity", "Govt Bond", "Corp Bond", "Cash", "Commodity", "FX Forward"]
HAIRCUTS = {
    "Equity": 0.15,
    "Govt Bond": 0.02,
    "Corp Bond": 0.08,
    "Cash": 0.0,
    "Commodity": 0.2,
    "FX Forward": 0.12,
}
LIQUIDITY_BANDS = {
    "Equity": "L2",
    "Govt Bond": "L1",
    "Corp Bond": "L2",
    "Cash": "L1",
    "Commodity": "L3",
    "FX Forward": "L2",
}
RATINGS = ["AAA", "AA", "A", "BBB", "BB", "B"]
RISK_WEIGHTS = {"AAA": 0.2, "AA": 0.3, "A": 0.5, "BBB": 1.0, "BB": 1.5, "B": 2.0}
CURRENCIES = ["GBP", "USD", "EUR"]

CLEAN_ROWS = 500
CAPITAL_RATIO = 0.105


def build_clean_pipeline(path: Path) -> None:
    """The well-behaved case, and the R1C1 compression showcase.

    Eight dense whole-column formula fills over 500 rows collapse to eight
    logical operations. This is also the only fixture carrying cached
    calculated values, so it is the positive path for reconciliation.
    """
    rng = random.Random(SEED)
    workbook = new_workbook()

    # -- Data ---------------------------------------------------------------
    data = workbook.active
    data.title = "Data"
    write_header(
        data,
        1,
        ["trade_id", "counterparty", "asset_class", "notional", "currency", "trade_date", "rating"],
    )
    rows: list[dict[str, Any]] = []
    base_date = dt.date(2026, 1, 5)
    for index in range(CLEAN_ROWS):
        record = {
            "trade_id": f"TRD-{index + 1:05d}",
            "counterparty": rng.choice(COUNTERPARTIES),
            "asset_class": rng.choice(ASSET_CLASSES),
            "notional": round(rng.uniform(10_000, 5_000_000), 2),
            "currency": rng.choice(CURRENCIES),
            "trade_date": base_date + dt.timedelta(days=rng.randint(0, 175)),
            "rating": rng.choice(RATINGS),
        }
        rows.append(record)
        row = index + 2
        data.cell(row=row, column=1, value=record["trade_id"])
        data.cell(row=row, column=2, value=record["counterparty"])
        data.cell(row=row, column=3, value=record["asset_class"])
        data.cell(row=row, column=4, value=record["notional"])
        data.cell(row=row, column=5, value=record["currency"])
        date_cell = data.cell(row=row, column=6, value=record["trade_date"])
        date_cell.number_format = "yyyy-mm-dd"
        data.cell(row=row, column=7, value=record["rating"])
    data.freeze_panes = "A2"

    # -- Ref ----------------------------------------------------------------
    ref = workbook.create_sheet("Ref")
    write_header(ref, 1, ["asset_class", "haircut", "liquidity_band"])
    for index, asset_class in enumerate(ASSET_CLASSES):
        row = index + 2
        ref.cell(row=row, column=1, value=asset_class)
        ref.cell(row=row, column=2, value=HAIRCUTS[asset_class])
        ref.cell(row=row, column=3, value=LIQUIDITY_BANDS[asset_class])

    # -- Params -------------------------------------------------------------
    params = workbook.create_sheet("Params")
    write_header(params, 1, ["parameter", "value"])
    params["A2"], params["B2"] = "capital_ratio", CAPITAL_RATIO
    params["A3"], params["B3"] = "fx_rate_usd", 0.79
    params["A4"], params["B4"] = "reporting_date", dt.date(2026, 6, 30)
    params["B4"].number_format = "yyyy-mm-dd"
    params["A5"], params["B5"] = "tolerance", 0.01
    for name, ref_text in [
        ("capital_ratio", "Params!$B$2"),
        ("fx_rate_usd", "Params!$B$3"),
        ("reporting_date", "Params!$B$4"),
        ("tolerance", "Params!$B$5"),
    ]:
        workbook.defined_names[name] = DefinedName(name, attr_text=ref_text)

    # -- Calc ---------------------------------------------------------------
    calc = workbook.create_sheet("Calc")
    write_header(
        calc,
        1,
        [
            "trade_id",
            "asset_class",
            "notional",
            "haircut",
            "adjusted_notional",
            "risk_weight",
            "rwa",
            "capital_charge",
        ],
    )
    cached_calc: dict[str, Any] = {}
    last = CLEAN_ROWS + 1
    for index, record in enumerate(rows):
        row = index + 2
        calc[f"A{row}"] = f"=Data!A{row}"
        calc[f"B{row}"] = f"=Data!C{row}"
        calc[f"C{row}"] = f"=Data!D{row}"
        calc[f"D{row}"] = f"=VLOOKUP(B{row},Ref!$A$2:$C$7,2,FALSE)"
        calc[f"E{row}"] = f"=C{row}*(1-D{row})"
        calc[f"F{row}"] = (
            f'=IF(Data!G{row}="AAA",0.2,IF(Data!G{row}="AA",0.3,'
            f'IF(Data!G{row}="A",0.5,IF(Data!G{row}="BBB",1,'
            f'IF(Data!G{row}="BB",1.5,2)))))'
        )
        calc[f"G{row}"] = f"=ROUND(E{row}*F{row},2)"
        calc[f"H{row}"] = f"=ROUND(G{row}*capital_ratio,2)"

        haircut = HAIRCUTS[record["asset_class"]]
        adjusted = record["notional"] * (1 - haircut)
        weight = RISK_WEIGHTS[record["rating"]]
        rwa = excel_round(adjusted * weight, 2)
        charge = excel_round(rwa * CAPITAL_RATIO, 2)
        cached_calc.update(
            {
                f"A{row}": record["trade_id"],
                f"B{row}": record["asset_class"],
                f"C{row}": record["notional"],
                f"D{row}": haircut,
                f"E{row}": adjusted,
                f"F{row}": weight,
                f"G{row}": rwa,
                f"H{row}": charge,
            }
        )
        record["rwa"] = rwa
        record["capital_charge"] = charge
        record["adjusted"] = adjusted
        record["haircut"] = haircut

    # -- Output -------------------------------------------------------------
    output = workbook.create_sheet("Output")
    write_header(output, 1, ["metric", "value"])
    metrics = [
        (
            "total_notional",
            f"=ROUND(SUM(Calc!C2:C{last}),2)",
            excel_round(sum(r["notional"] for r in rows), 2),
        ),
        (
            "total_rwa",
            f"=ROUND(SUM(Calc!G2:G{last}),2)",
            excel_round(sum(r["rwa"] for r in rows), 2),
        ),
        (
            "total_capital",
            f"=ROUND(SUM(Calc!H2:H{last}),2)",
            excel_round(sum(r["capital_charge"] for r in rows), 2),
        ),
        ("trade_count", f"=COUNT(Calc!C2:C{last})", len(rows)),
        (
            "equity_rwa",
            f'=ROUND(SUMIF(Calc!B2:B{last},"Equity",Calc!G2:G{last}),2)',
            excel_round(sum(r["rwa"] for r in rows if r["asset_class"] == "Equity"), 2),
        ),
        (
            "mean_haircut",
            f"=ROUND(AVERAGE(Calc!D2:D{last}),4)",
            excel_round(sum(r["haircut"] for r in rows) / len(rows), 4),
        ),
    ]
    cached_output: dict[str, Any] = {}
    for index, (name, formula, value) in enumerate(metrics):
        row = index + 2
        output.cell(row=row, column=1, value=name)
        output.cell(row=row, column=2, value=formula)
        cached_output[f"B{row}"] = value

    for sheet, width in [(data, 16), (calc, 18), (output, 18), (ref, 16), (params, 18)]:
        sheet.column_dimensions["A"].width = width

    workbook.save(path)

    # -- post-process: inject the cached values -----------------------------
    parts = read_parts(path)
    inject_cached_values(parts, {"Calc": cached_calc, "Output": cached_output})
    set_full_calc_on_load(parts, enabled=False)
    write_parts(path, parts)


# =============================================================================
# Fixture 2: powerquery.xlsx
# =============================================================================

SECTION_1_M = """section Section1;

shared Exposures = let
    Source = Excel.CurrentWorkbook(){[Name="tbl_Exposures"]}[Content],
    #"Changed Type" = Table.TransformColumnTypes(Source,{{"trade_id", type text}, {"counterparty", type text}, {"notional", type number}, {"asset_class", type text}}),
    #"Filtered Rows" = Table.SelectRows(#"Changed Type", each [notional] > 0),
    #"Added Custom" = Table.AddColumn(#"Filtered Rows", "notional_gbp", each [notional] * 0.79, type number)
in
    #"Added Custom";

shared CollateralHaircuts = let
    Source = Excel.CurrentWorkbook(){[Name="tbl_Haircuts"]}[Content],
    #"Changed Type" = Table.TransformColumnTypes(Source,{{"asset_class", type text}, {"haircut", type number}})
in
    #"Changed Type";

shared NetExposure = let
    Source = Table.NestedJoin(Exposures, {"asset_class"}, CollateralHaircuts, {"asset_class"}, "haircut_row", JoinKind.LeftOuter),
    #"Expanded haircut_row" = Table.ExpandTableColumn(Source, "haircut_row", {"haircut"}, {"haircut"}),
    #"Replaced Errors" = Table.ReplaceErrorValues(#"Expanded haircut_row", {{"haircut", 0}}),
    #"Added Net" = Table.AddColumn(#"Replaced Errors", "net_exposure", each [notional_gbp] * (1 - [haircut]), type number),
    #"Removed Columns" = Table.RemoveColumns(#"Added Net", {"haircut_row"})
in
    #"Removed Columns";
"""

POWER_QUERY_NAMES = ["Exposures", "CollateralHaircuts", "NetExposure"]

DECOY_ITEM_1 = (
    b'<?xml version="1.0" encoding="utf-8"?>'
    b'<properties xmlns="http://schemas.microsoft.com/office/2006/metadata/properties">'
    b"<documentManagement/>"
    b"</properties>"
)

DECOY_ITEM_2 = (
    b'<?xml version="1.0" encoding="utf-8"?>'
    b'<ReportMetadata xmlns="urn:acme-bank:reporting:v2">'
    b"<Owner>Market Risk</Owner>"
    b"<Classification>Internal</Classification>"
    b"<ReviewCycle>Monthly</ReviewCycle>"
    b"</ReportMetadata>"
)


def build_powerquery(path: Path) -> None:
    """A real Power Query DataMashup, deliberately at item3.xml behind two decoys.

    PLAN section 1.5 warns against hardcoding ``customXml/item1.xml``. The two
    decoy parts are valid custom XML with unrelated schemas, so an extractor
    must iterate and match on the DataMashup namespace to find the right one.
    """
    workbook = new_workbook()

    report = workbook.active
    report.title = "Report"
    write_header(report, 1, ["asset_class", "net_exposure", "share"])
    exposures = [
        ("Equity", 4_820_331.44),
        ("Govt Bond", 12_004_918.02),
        ("Corp Bond", 7_331_202.85),
        ("Cash", 2_118_440.00),
        ("Commodity", 964_775.19),
        ("FX Forward", 3_402_664.71),
    ]
    total = sum(value for _, value in exposures)
    for index, (asset_class, value) in enumerate(exposures):
        row = index + 2
        report.cell(row=row, column=1, value=asset_class)
        report.cell(row=row, column=2, value=value)
        report.cell(row=row, column=3, value=f"=ROUND(B{row}/${'B'}${len(exposures) + 3},4)")
    report.cell(row=len(exposures) + 3, column=1, value="total")
    report.cell(row=len(exposures) + 3, column=2, value=f"=SUM(B2:B{len(exposures) + 1})")
    assert abs(total - sum(v for _, v in exposures)) < 1e-9

    staging = workbook.create_sheet("Staging")
    write_header(staging, 1, ["trade_id", "counterparty", "notional", "asset_class"])
    rng = random.Random(SEED + 2)
    for index in range(40):
        row = index + 2
        staging.cell(row=row, column=1, value=f"TRD-{index + 1:05d}")
        staging.cell(row=row, column=2, value=rng.choice(COUNTERPARTIES))
        staging.cell(row=row, column=3, value=round(rng.uniform(5_000, 900_000), 2))
        staging.cell(row=row, column=4, value=rng.choice(ASSET_CLASSES))

    haircuts = workbook.create_sheet("HaircutTable")
    write_header(haircuts, 1, ["asset_class", "haircut"])
    for index, asset_class in enumerate(ASSET_CLASSES):
        haircuts.cell(row=index + 2, column=1, value=asset_class)
        haircuts.cell(row=index + 2, column=2, value=HAIRCUTS[asset_class])

    workbook.save(path)

    parts = read_parts(path)
    envelope = build_datamashup(SECTION_1_M, POWER_QUERY_NAMES)
    mashup_item = (
        '<?xml version="1.0" encoding="utf-8"?>'
        f'<DataMashup sqmid="00000000-0000-0000-0000-000000000000" xmlns="{DATAMASHUP_NS}">'
        f"{base64.b64encode(envelope).decode('ascii')}"
        "</DataMashup>"
    ).encode()

    attach_custom_xml(
        parts,
        [
            (
                DECOY_ITEM_1,
                "http://schemas.microsoft.com/office/2006/metadata/properties",
                "{6C1E1A5C-0000-4E3C-9F1B-000000000001}",
            ),
            (DECOY_ITEM_2, "urn:acme-bank:reporting:v2", "{6C1E1A5C-0000-4E3C-9F1B-000000000002}"),
            (mashup_item, DATAMASHUP_NS, "{6C1E1A5C-0000-4E3C-9F1B-000000000003}"),
        ],
    )
    write_parts(path, parts)


# =============================================================================
# Fixture 3: legacy_sql.xlsx
# =============================================================================

ODBC_SQL = """SELECT
    t.trade_id,
    t.counterparty_id,
    c.counterparty_name,
    t.notional,
    t.currency,
    t.trade_date,
    r.rating_grade
FROM risk.trades AS t
INNER JOIN risk.counterparties AS c
    ON c.counterparty_id = t.counterparty_id
LEFT JOIN risk.ratings AS r
    ON r.counterparty_id = t.counterparty_id
   AND r.as_of_date = t.trade_date
WHERE t.trade_date >= ?
  AND t.book IN ('TRADING', 'BANKING')
  AND t.status <> 'CANCELLED'
ORDER BY t.trade_date DESC"""

OLEDB_SQL = """SELECT
    period_end,
    legal_entity,
    SUM(rwa_amount) AS total_rwa
FROM finance.rwa_summary
WHERE period_end = ?
GROUP BY period_end, legal_entity"""

ODBC_CONNECTION_STRING = (
    "ODBC;DSN=RiskWarehouse;Description=Risk Warehouse (PROD);UID=svc_riskread;"
    "Trusted_Connection=Yes;APP=Microsoft Office 2016;WSID=LDN-RISK-042;"
    "DATABASE=RiskWarehouse"
)
OLEDB_CONNECTION_STRING = (
    "OLEDB;Provider=MSOLAP.8;Integrated Security=SSPI;Persist Security Info=True;"
    "Initial Catalog=FinanceCube;Data Source=ldn-olap-01;MDX Compatibility=1;"
    "Safety Options=2;MDX Missing Member Mode=Error;Update Isolation Level=2"
)


def build_legacy_sql(path: Path) -> None:
    """Two legacy connections carrying multi-line SQL, one ODBC and one OLEDB."""
    workbook = new_workbook()

    extract = workbook.active
    extract.title = "Extract"
    write_header(
        extract,
        1,
        ["trade_id", "counterparty_name", "notional", "currency", "trade_date", "rating_grade"],
    )
    rng = random.Random(SEED + 3)
    for index in range(120):
        row = index + 2
        extract.cell(row=row, column=1, value=f"TRD-{index + 1:05d}")
        extract.cell(row=row, column=2, value=rng.choice(COUNTERPARTIES))
        extract.cell(row=row, column=3, value=round(rng.uniform(20_000, 3_000_000), 2))
        extract.cell(row=row, column=4, value=rng.choice(CURRENCIES))
        date_cell = extract.cell(
            row=row, column=5, value=dt.date(2026, 5, 1) + dt.timedelta(days=rng.randint(0, 60))
        )
        date_cell.number_format = "yyyy-mm-dd"
        extract.cell(row=row, column=6, value=rng.choice(RATINGS))

    summary = workbook.create_sheet("Summary")
    write_header(summary, 1, ["currency", "total_notional", "trade_count"])
    for index, currency in enumerate(CURRENCIES):
        row = index + 2
        summary.cell(row=row, column=1, value=currency)
        summary.cell(
            row=row, column=2, value=f"=SUMIF(Extract!$D$2:$D$121,A{row},Extract!$C$2:$C$121)"
        )
        summary.cell(row=row, column=3, value=f"=COUNTIF(Extract!$D$2:$D$121,A{row})")

    workbook.save(path)

    connections_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<connections xmlns="{SHEET_MAIN_NS}">'
        '<connection id="1" name="RiskWarehouse" type="1" refreshedVersion="8"'
        ' minRefreshableVersion="3" background="1" saveData="1"'
        ' description="Trade extract from the risk warehouse (nightly)">'
        f'<dbPr connection="{xml_attr(ODBC_CONNECTION_STRING)}"'
        f' command="{xml_attr(ODBC_SQL)}" commandType="2"/>'
        "</connection>"
        '<connection id="2" name="FinanceCube" type="5" refreshedVersion="8"'
        ' minRefreshableVersion="3" background="1" saveData="1"'
        ' description="Finance RWA summary cube">'
        f'<dbPr connection="{xml_attr(OLEDB_CONNECTION_STRING)}"'
        f' command="{xml_attr(OLEDB_SQL)}" commandType="2"/>'
        "</connection>"
        "</connections>"
    ).encode()

    parts = read_parts(path)
    attach_connections(parts, connections_xml)
    write_parts(path, parts)


# =============================================================================
# Fixture 4: cross_sheet_chain.xlsx
# =============================================================================

CHAIN_ROWS = 120


def build_cross_sheet_chain(path: Path) -> None:
    """A four-deep cross-sheet dependency chain for the DAG builder.

    A_Source to B_Enrich to C_Aggregate to D_Report, with absolute parameter
    references, INDEX/MATCH, SUMIFS, SUMPRODUCT, a running total and a
    row-above reference.
    """
    rng = random.Random(SEED + 4)
    workbook = new_workbook()

    params = workbook.active
    params.title = "Params"
    params["A1"] = "uplift_factor"
    params["B1"] = 1.075
    params["A2"] = "floor_amount"
    params["B2"] = 250.0
    params["A3"] = "reporting_ccy"
    params["B3"] = "GBP"
    workbook.defined_names["uplift_factor"] = DefinedName("uplift_factor", attr_text="Params!$B$1")

    source = workbook.create_sheet("A_Source")
    write_header(source, 1, ["book", "region", "gross_amount", "units"])
    books = ["TRADING", "BANKING", "TREASURY"]
    regions = ["EMEA", "AMER", "APAC"]
    records: list[tuple[str, str, float, int]] = []
    for index in range(CHAIN_ROWS):
        row = index + 2
        record = (
            rng.choice(books),
            rng.choice(regions),
            round(rng.uniform(100, 90_000), 2),
            rng.randint(1, 400),
        )
        records.append(record)
        for column, value in enumerate(record, start=1):
            source.cell(row=row, column=column, value=value)

    last = CHAIN_ROWS + 1

    # B depends on A and on an absolute parameter reference.
    enrich = workbook.create_sheet("B_Enrich")
    write_header(enrich, 1, ["book", "region", "uplifted", "unit_price", "running_total"])
    for index in range(CHAIN_ROWS):
        row = index + 2
        enrich[f"A{row}"] = f"=A_Source!A{row}"
        enrich[f"B{row}"] = f"=A_Source!B{row}"
        enrich[f"C{row}"] = f"=A_Source!C{row}*Params!$B$1"
        enrich[f"D{row}"] = f"=IFERROR(C{row}/A_Source!D{row},Params!$B$2)"
        # Running total: the classic cum_sum shape, and a row-above reference.
        enrich[f"E{row}"] = f"=C{row}" if row == 2 else f"=E{row - 1}+C{row}"

    # C aggregates B with SUMIFS and INDEX/MATCH.
    aggregate = workbook.create_sheet("C_Aggregate")
    write_header(aggregate, 1, ["book", "region", "total", "weighted", "first_price"])
    combinations = [(book, region) for book in books for region in regions]
    for index, (book, region) in enumerate(combinations):
        row = index + 2
        aggregate.cell(row=row, column=1, value=book)
        aggregate.cell(row=row, column=2, value=region)
        aggregate[f"C{row}"] = (
            f"=SUMIFS(B_Enrich!$C$2:$C${last},B_Enrich!$A$2:$A${last},$A{row},"
            f"B_Enrich!$B$2:$B${last},$B{row})"
        )
        aggregate[f"D{row}"] = (
            f"=SUMPRODUCT((B_Enrich!$A$2:$A${last}=$A{row})*"
            f"(B_Enrich!$B$2:$B${last}=$B{row}),B_Enrich!$C$2:$C${last})"
        )
        aggregate[f"E{row}"] = (
            f"=IFERROR(INDEX(B_Enrich!$D$2:$D${last},MATCH($A{row},B_Enrich!$A$2:$A${last},0)),0)"
        )

    # D reports over C, closing the four-level chain.
    report = workbook.create_sheet("D_Report")
    write_header(report, 1, ["metric", "value"])
    combo_last = len(combinations) + 1
    report["A2"], report["B2"] = "grand_total", f"=SUM(C_Aggregate!C2:C{combo_last})"
    report["A3"], report["B3"] = "weighted_total", f"=SUM(C_Aggregate!D2:D{combo_last})"
    report["A4"], report["B4"] = "reconciliation_delta", "=ROUND(B2-B3,6)"
    report["A5"], report["B5"] = "uplift_applied", "=uplift_factor"
    report["A6"], report["B6"] = "source_total", f"=SUM(A_Source!C2:C{last})"
    report["A7"], report["B7"] = "implied_uplift", "=IFERROR(ROUND(B2/B6,6),0)"
    report["A8"], report["B8"] = (
        "emea_share",
        (
            f'=IFERROR(ROUND(SUMIF(C_Aggregate!$B$2:$B${combo_last},"EMEA",'
            f"C_Aggregate!$C$2:$C${combo_last})/B2,4),0)"
        ),
    )
    report["A9"], report["B9"] = "chain_depth_probe", f"=B_Enrich!E{last}"

    save(workbook, path)


# =============================================================================
# Fixture 5: mostly_manual.xlsx
# =============================================================================


def build_mostly_manual(path: Path) -> None:
    """Mostly typed values with a thin calculation layer.

    Sets two traps: a block of hardcoded overrides with no formula at all (the
    checkpoint-stage trigger from PLAN section 2.2), and magic constants
    embedded directly in formulas where a parameter reference belonged.
    """
    rng = random.Random(SEED + 5)
    workbook = new_workbook()

    entry = workbook.active
    entry.title = "Input"
    write_header(entry, 1, ["cost_centre", "description", "amount", "approved_by"])
    centres = ["CC-1100", "CC-1200", "CC-2300", "CC-3100", "CC-4200"]
    descriptions = [
        "Market data licence",
        "Contractor day rate",
        "Software renewal",
        "Travel and expenses",
        "Hardware refresh",
        "Training budget",
        "Consultancy fee",
        "Office services",
    ]
    approvers = ["j.okafor", "m.lindqvist", "s.chaudhry", "p.moreau"]
    for index in range(60):
        row = index + 2
        entry.cell(row=row, column=1, value=rng.choice(centres))
        entry.cell(row=row, column=2, value=rng.choice(descriptions))
        entry.cell(row=row, column=3, value=round(rng.uniform(500, 48_000), 2))
        entry.cell(row=row, column=4, value=rng.choice(approvers))

    # The overrides block: typed numbers, no formula, no provenance. This is the
    # shape that must become a checkpoint stage rather than generated code.
    overrides = workbook.create_sheet("Overrides")
    overrides["A1"] = "Manual overrides agreed with Finance — do not automate"
    overrides["A1"].font = Font(bold=True, size=12)
    write_header(overrides, 3, ["cost_centre", "override_amount", "reason", "agreed_date"])
    override_rows = [
        ("CC-1100", 12_500.00, "Accrual reversal agreed with FinCon", dt.date(2026, 6, 12)),
        ("CC-1200", -3_400.00, "Duplicate invoice removed", dt.date(2026, 6, 12)),
        ("CC-2300", 8_750.50, "Late licence true-up", dt.date(2026, 6, 15)),
        ("CC-3100", 1_200.00, "Recharge from Ops", dt.date(2026, 6, 15)),
        ("CC-4200", -975.25, "Cancelled training", dt.date(2026, 6, 18)),
        ("CC-1100", 4_000.00, "Q2 catch-up", dt.date(2026, 6, 20)),
        ("CC-2300", 620.00, "Rounding to agreed total", dt.date(2026, 6, 20)),
        ("CC-3100", -150.00, "Reversal", dt.date(2026, 6, 21)),
        ("CC-1200", 2_310.75, "Contractor extension", dt.date(2026, 6, 22)),
        ("CC-4200", 500.00, "Ad-hoc spend", dt.date(2026, 6, 22)),
        ("CC-1100", -80.00, "Immaterial correction", dt.date(2026, 6, 23)),
        ("CC-2300", 15_000.00, "Strategic project uplift", dt.date(2026, 6, 24)),
        ("CC-3100", 240.00, "Stationery recharge", dt.date(2026, 6, 24)),
        ("CC-4200", -1_100.00, "Budget reallocation", dt.date(2026, 6, 25)),
    ]
    for index, (centre, amount, reason, agreed) in enumerate(override_rows):
        row = index + 4
        overrides.cell(row=row, column=1, value=centre)
        overrides.cell(row=row, column=2, value=amount)
        overrides.cell(row=row, column=3, value=reason)
        date_cell = overrides.cell(row=row, column=4, value=agreed)
        date_cell.number_format = "yyyy-mm-dd"

    # Thin calculation layer, riddled with magic constants.
    calc = workbook.create_sheet("Calc")
    write_header(calc, 1, ["cost_centre", "base", "vat", "contingency", "total"])
    for index, centre in enumerate(centres):
        row = index + 2
        calc.cell(row=row, column=1, value=centre)
        calc[f"B{row}"] = f"=SUMIF(Input!$A$2:$A$61,A{row},Input!$C$2:$C$61)"
        # 0.20 and 0.08 should have been parameters; both are Findings.
        calc[f"C{row}"] = f"=ROUND(B{row}*0.2,2)"
        calc[f"D{row}"] = f"=ROUND(B{row}*0.08,2)"
        calc[f"E{row}"] = (
            f"=B{row}+C{row}+D{row}+SUMIF(Overrides!$A$4:$A$17,A{row},Overrides!$B$4:$B$17)"
        )

    output = workbook.create_sheet("Output")
    write_header(output, 1, ["metric", "value"])
    output["A2"], output["B2"] = "total_base", "=SUM(Calc!B2:B6)"
    output["A3"], output["B3"] = "total_with_overrides", "=SUM(Calc!E2:E6)"
    output["A4"], output["B4"] = "override_impact", "=ROUND(B3-B2,2)"
    # Another magic constant: a hardcoded headcount divisor.
    output["A5"], output["B5"] = "cost_per_head", "=ROUND(B3/47,2)"

    save(workbook, path)


# =============================================================================
# Fixture 6: documented.xlsx (+ Word companions)
# =============================================================================

PROCESS_NOTES = [
    (
        "Purpose",
        "This workbook produces the monthly counterparty exposure return submitted to the "
        "Risk Oversight committee on the fifth working day. It replaces the quarterly "
        "spreadsheet retired in 2023 and is owned by the Market Risk reporting team. The "
        "return covers all trading and banking book exposures above the de minimis "
        "threshold agreed with Finance.",
    ),
    (
        "Inputs",
        "The hand-in arrives from the Data Services team as a CSV extract dropped into the "
        "shared reporting folder, normally by close of business on the second working day. "
        "It must contain one row per trade with the trade identifier, counterparty, "
        "notional, currency and trade date. If the file has not arrived by 10:00 on the "
        "third working day, escalate to the Data Services duty manager before proceeding.",
    ),
    (
        "Step 1 — refresh reference data",
        "Open the Ref sheet and confirm the haircut table matches the current collateral "
        "policy document. The policy is reissued each January; if the version stamp in "
        "cell A1 of Ref does not match the current policy, do not proceed. Haircuts are "
        "applied by asset class only; there is no counterparty-specific override in this "
        "process and requests for one should be routed to the policy team.",
    ),
    (
        "Step 2 — load and reconcile",
        "Paste the hand-in into the Data sheet starting at cell A2. Do not disturb the "
        "header row. The Calc sheet formulas fill automatically. Check that the trade "
        "count on the Output sheet matches the control total supplied in the covering "
        "email from Data Services. A difference of more than five trades has always "
        "indicated a truncated extract rather than a genuine movement.",
    ),
    (
        "Step 3 — manual review",
        "Review any exposure above ten million for plausibility against last month's "
        "return. This is a judgement step and is not automated: the reviewer signs off in "
        "the sign-off box below and records anything unusual in the commentary. Historic "
        "practice has been to compare against the prior month rather than the prior "
        "quarter, because the book turns over quickly.",
    ),
    (
        "Known issues",
        "The VLOOKUP on the Calc sheet returns #N/A when a new asset class appears that is "
        "not yet in the reference table. When this happens, add the asset class to Ref and "
        "notify the policy team; do not simply wrap the lookup in IFERROR, because a "
        "silent zero haircut understates the exposure and this has caused a restatement "
        "before.",
    ),
    (
        "Sign-off",
        "Preparer and reviewer must both initial the sign-off box before the return is "
        "submitted. The submitted file is archived to the reporting share under the "
        "period end date. Retention is seven years under the records policy.",
    ),
]

DOCX_PARAGRAPHS = [
    ("Monthly Counterparty Exposure Return — Operating Procedure", "Title"),
    ("Scope", "Heading 1"),
    (
        "This procedure covers the preparation, review and submission of the monthly "
        "counterparty exposure return. It applies to the Market Risk reporting team and "
        "to any delegate covering the process during absence.",
        None,
    ),
    ("Roles", "Heading 1"),
    (
        "The preparer runs the workbook and completes the control checks. The reviewer "
        "performs the plausibility review and signs off. The process owner is the Head of "
        "Market Risk Reporting, who approves any change to the calculation basis.",
        None,
    ),
    ("Procedure", "Heading 1"),
    (
        "1. Confirm the hand-in has arrived and matches the control totals in the covering "
        "email. 2. Paste the extract into the Data sheet without disturbing the header "
        "row. 3. Confirm the reference haircut table matches the current collateral policy. "
        "4. Review exposures above ten million against the prior month. 5. Obtain reviewer "
        "sign-off. 6. Archive the submitted file to the reporting share.",
        None,
    ),
    ("Controls", "Heading 1"),
    (
        "The trade count control must agree to the figure supplied by Data Services. The "
        "reconciliation cell on the Output sheet must show a difference of zero before "
        "submission. Any manual adjustment must be recorded with a reason and the name of "
        "the person who agreed it.",
        None,
    ),
    ("Escalation", "Heading 1"),
    (
        "If the hand-in has not arrived by 10:00 on the third working day, escalate to the "
        "Data Services duty manager. If the reconciliation does not clear, escalate to the "
        "process owner before submitting anything.",
        None,
    ),
]


def build_documented(path: Path) -> None:
    """Prose process notes in a sheet, plus cell comments, for docs.py to find."""
    workbook = new_workbook()

    notes = workbook.active
    notes.title = "Process Notes"
    notes.column_dimensions["A"].width = 28
    notes.column_dimensions["B"].width = 110
    notes["A1"] = "Monthly Counterparty Exposure Return"
    notes["A1"].font = Font(bold=True, size=14)
    notes["A2"] = "Last reviewed"
    notes["B2"] = dt.date(2026, 6, 30)
    notes["B2"].number_format = "yyyy-mm-dd"
    for index, (heading, body) in enumerate(PROCESS_NOTES):
        row = index + 4
        heading_cell = notes.cell(row=row, column=1, value=heading)
        heading_cell.font = Font(bold=True)
        heading_cell.alignment = Alignment(vertical="top")
        body_cell = notes.cell(row=row, column=2, value=body)
        body_cell.alignment = Alignment(wrap_text=True, vertical="top")
        notes.row_dimensions[row].height = 60

    data = workbook.create_sheet("Data")
    write_header(data, 1, ["trade_id", "counterparty", "asset_class", "notional"])
    rng = random.Random(SEED + 6)
    for index in range(80):
        row = index + 2
        data.cell(row=row, column=1, value=f"TRD-{index + 1:05d}")
        data.cell(row=row, column=2, value=rng.choice(COUNTERPARTIES))
        data.cell(row=row, column=3, value=rng.choice(ASSET_CLASSES))
        data.cell(row=row, column=4, value=round(rng.uniform(1_000, 12_000_000), 2))
    data["A1"].comment = Comment(
        "Trade identifier as supplied by Data Services. Do not re-key; a mismatch here "
        "breaks the join to the prior month file.",
        "m.lindqvist",
        height=110,
        width=280,
    )
    data["D1"].comment = Comment(
        "Notional is always in trade currency, never reporting currency. The FX "
        "conversion happens on the Calc sheet.",
        "j.okafor",
        height=110,
        width=280,
    )

    calc = workbook.create_sheet("Calc")
    write_header(calc, 1, ["trade_id", "notional", "haircut", "net"])
    for index in range(80):
        row = index + 2
        calc[f"A{row}"] = f"=Data!A{row}"
        calc[f"B{row}"] = f"=Data!D{row}"
        calc[f"C{row}"] = f"=VLOOKUP(Data!C{row},Ref!$A$2:$B$7,2,FALSE)"
        calc[f"D{row}"] = f"=ROUND(B{row}*(1-C{row}),2)"
    calc["C1"].comment = Comment(
        "Returns #N/A for an unmapped asset class. This is deliberate — see the Known "
        "issues section of the Process Notes sheet. Do not wrap in IFERROR.",
        "m.lindqvist",
        height=120,
        width=300,
    )

    ref = workbook.create_sheet("Ref")
    write_header(ref, 1, ["asset_class", "haircut"])
    for index, asset_class in enumerate(ASSET_CLASSES):
        ref.cell(row=index + 2, column=1, value=asset_class)
        ref.cell(row=index + 2, column=2, value=HAIRCUTS[asset_class])

    save(workbook, path)


def build_word_companions(directory: Path) -> None:
    """Write the .docx procedure and the deliberately fake .doc stub."""
    from docx import Document  # imported lazily; only this fixture needs it

    document = Document()
    for text, style in DOCX_PARAGRAPHS:
        document.add_paragraph(text, style=style) if style else document.add_paragraph(text)
    document.core_properties.created = FIXED_TIMESTAMP
    document.core_properties.modified = FIXED_TIMESTAMP
    document.core_properties.author = "kedge fixture generator"
    document.core_properties.last_modified_by = "kedge fixture generator"
    document.core_properties.title = "Monthly Counterparty Exposure Return — Operating Procedure"
    docx_path = directory / "documented_procedure.docx"
    document.save(docx_path)
    normalise(docx_path)

    # Not a real OLE2 compound document: python-docx must fail with a clear
    # conversion hint rather than a confusing zip or XML error (PLAN section 1.5).
    stub = (
        "This file stands in for a genuine Word 97-2003 binary .doc.\r\n"
        "It is intentionally NOT an OLE2 compound file, so any attempt to parse it "
        "must fail clearly and suggest converting to .docx.\r\n"
    )
    (directory / "procedure_legacy.doc").write_bytes(stub.encode("utf-8"))


# =============================================================================
# Fixture 7: hostile.xlsx
# =============================================================================

HOSTILE_DATA_FIRST_ROW = 5
HOSTILE_DATA_LAST_ROW = 80
HOSTILE_SUBTOTAL_ROW = 30
HOSTILE_INCONSISTENT_ROW = 47
HOSTILE_BLANK_ROWS = (45, 60)
HOSTILE_GRAND_TOTAL_ROW = 82


def build_hostile(path: Path) -> None:
    """Deliberately awful. The analyser must survive this and produce findings.

    Every trap is listed in README.md with the row or cell it occupies, so a
    corpus test can assert on specific coordinates rather than counts alone.
    """
    rng = random.Random(SEED + 7)
    workbook = new_workbook()

    messy = workbook.active
    messy.title = "Messy"

    # Three preamble rows before the real header, with a merged title.
    messy.merge_cells("A1:F1")
    messy["A1"] = "ACME BANK — MONTHLY POSITION EXTRACT (DO NOT CIRCULATE)"
    messy["A1"].font = Font(bold=True, size=14)
    messy["A1"].alignment = Alignment(horizontal="center")
    messy["A2"] = "Produced by:"
    messy["B2"] = "s.chaudhry"
    messy["A3"] = "Period:"
    messy["B3"] = "Jun-26"

    # Duplicate header names ("Amount" twice), and a blank column at F.
    write_header(
        messy,
        4,
        ["Ref", "Amount", "Rate", "Amount", "Notes", "", "Booked", "Legacy Code", "Spare"],
    )

    text_dates = ["01/03/2026", "2026-03-04", "7 Mar 2026", "Mar-26", "03/15/2026"]
    for index in range(HOSTILE_DATA_FIRST_ROW, HOSTILE_DATA_LAST_ROW + 1):
        if index in HOSTILE_BLANK_ROWS or index == HOSTILE_SUBTOTAL_ROW:
            continue
        messy.cell(row=index, column=1, value=f"R-{index:04d}")
        messy.cell(row=index, column=2, value=round(rng.uniform(100, 9_000), 2))
        messy.cell(row=index, column=3, value=round(rng.uniform(0.5, 1.5), 4))
        # Column D is the uniform region that row 47 breaks.
        messy.cell(row=index, column=4, value=f"=B{index}*C{index}")
        # Column E: numbers stored as text.
        messy.cell(row=index, column=5, value=f"{rng.uniform(10, 999):.2f}")
        # Column G: dates as text in mixed formats.
        messy.cell(row=index, column=7, value=text_dates[index % len(text_dates)])
        messy.cell(row=index, column=8, value=rng.choice(["LGC-A", "LGC-B", "LGC-C"]))

    # The classic "someone typed over row 47" bug: an inconsistent formula
    # inside an otherwise uniform R1C1 region.
    messy[f"D{HOSTILE_INCONSISTENT_ROW}"] = (
        f"=B{HOSTILE_INCONSISTENT_ROW}*C{HOSTILE_INCONSISTENT_ROW}*1.1"
    )

    # A totals row in the MIDDLE of the data.
    messy.cell(row=HOSTILE_SUBTOTAL_ROW, column=1, value="Subtotal to date")
    messy.cell(row=HOSTILE_SUBTOTAL_ROW, column=1).font = Font(bold=True)
    messy[f"B{HOSTILE_SUBTOTAL_ROW}"] = (
        f"=SUM(B{HOSTILE_DATA_FIRST_ROW}:B{HOSTILE_SUBTOTAL_ROW - 1})"
    )
    messy[f"D{HOSTILE_SUBTOTAL_ROW}"] = (
        f"=SUM(D{HOSTILE_DATA_FIRST_ROW}:D{HOSTILE_SUBTOTAL_ROW - 1})"
    )

    # ...and another at the bottom, after a gap.
    messy.cell(row=HOSTILE_GRAND_TOTAL_ROW, column=1, value="GRAND TOTAL")
    messy.cell(row=HOSTILE_GRAND_TOTAL_ROW, column=1).font = Font(bold=True)
    messy[f"B{HOSTILE_GRAND_TOTAL_ROW}"] = (
        f"=SUM(B{HOSTILE_DATA_FIRST_ROW}:B{HOSTILE_DATA_LAST_ROW})-B{HOSTILE_SUBTOTAL_ROW}"
    )
    messy[f"D{HOSTILE_GRAND_TOTAL_ROW}"] = (
        f"=SUM(D{HOSTILE_DATA_FIRST_ROW}:D{HOSTILE_DATA_LAST_ROW})-D{HOSTILE_SUBTOTAL_ROW}"
    )

    # A merged cell inside the data region, which breaks naive row reads.
    messy.merge_cells(f"A{HOSTILE_BLANK_ROWS[1] + 1}:B{HOSTILE_BLANK_ROWS[1] + 1}")
    messy.cell(row=HOSTILE_BLANK_ROWS[1] + 1, column=1, value="continued below")

    # Hidden columns.
    messy.column_dimensions["H"].hidden = True
    messy.column_dimensions["I"].hidden = True

    # IFERROR swallowing a genuine division by zero.
    messy["K5"] = "swallowed_error"
    messy["L5"] = 0
    messy["K6"] = "=IFERROR(B5/L5,0)"
    messy["K7"] = '=IFERROR(VLOOKUP("NOT-A-KEY",A5:B80,2,FALSE),"")'

    # A hard #REF! error cell, and a formula referring to a deleted range.
    messy["K9"] = "#REF!"
    messy["K10"] = "=SUM(#REF!)"

    # -- Volatile ------------------------------------------------------------
    volatile = workbook.create_sheet("Volatile")
    write_header(volatile, 1, ["what", "formula"])
    volatile["A2"], volatile["B2"] = "now", "=NOW()"
    volatile["A3"], volatile["B3"] = "today", "=TODAY()"
    volatile["A4"], volatile["B4"] = "random", "=RAND()"
    volatile["A5"], volatile["B5"] = "random_between", "=RANDBETWEEN(1,100)"
    volatile["A6"], volatile["B6"] = "offset", "=SUM(OFFSET(Messy!$B$5,0,0,10,1))"
    volatile["A7"], volatile["B7"] = "indirect", '=INDIRECT("Messy!B"&ROW())'
    volatile["A8"], volatile["B8"] = "indirect_sheet", '=INDIRECT("\'"&$A$9&"\'!B5")'
    volatile["A9"] = "Messy"

    # -- Circular ------------------------------------------------------------
    circular = workbook.create_sheet("Circular")
    circular["A1"] = "Genuine three-cell cycle: C2 -> C4 -> C3 -> C2"
    circular["C2"] = "=C4*2"
    circular["C3"] = "=C2+1"
    circular["C4"] = "=C3-5"
    circular["E2"] = "Direct self-reference"
    circular["E3"] = "=E3+1"

    # -- Broken --------------------------------------------------------------
    broken = workbook.create_sheet("Broken")
    broken["A1"] = "External workbook link that cannot resolve"
    broken["A2"] = "=[1]Rates!$B$2"
    broken["A3"] = "=[1]Rates!$B$3*2"
    broken["A4"] = "Defined name pointing at a deleted range"
    broken["A5"] = "=SUM(obsolete_rate_table)"

    # -- Hidden sheet --------------------------------------------------------
    archive = workbook.create_sheet("_Archive")
    write_header(archive, 1, ["period", "value"])
    for index in range(12):
        archive.cell(row=index + 2, column=1, value=f"2025-{index + 1:02d}")
        archive.cell(row=index + 2, column=2, value=round(rng.uniform(1000, 5000), 2))
    archive.sheet_state = "hidden"

    # A defined name pointing at a range that no longer exists.
    workbook.defined_names["obsolete_rate_table"] = DefinedName(
        "obsolete_rate_table", attr_text="#REF!#REF!"
    )
    workbook.defined_names["messy_amounts"] = DefinedName(
        "messy_amounts",
        attr_text=f"Messy!$B${HOSTILE_DATA_FIRST_ROW}:$B${HOSTILE_DATA_LAST_ROW}",
    )

    # The unresolvable external workbook, written as a real externalLink part.
    link = ExternalLink(
        externalBook=ExternalBook(sheetNames=ExternalSheetNames(sheetName=["Rates"]))
    )
    link.file_link = Relationship(
        Id="rId1",
        Type=f"{OFFICE_REL}/externalLinkPath",
        Target="file:///Z:/retired/vendor_rates_2019.xlsx",
        TargetMode="External",
    )
    workbook._external_links.append(link)

    save(workbook, path)


# =============================================================================
# Fixture 8: no_cached_values.xlsx
# =============================================================================


def build_no_cached_values(path: Path) -> None:
    """A tool-written workbook: every formula cell reads None under data_only.

    This is the PLAN section 1.5 caveat. Reconciliation must be disabled for
    this workbook, not silently reported as zeros.
    """
    rng = random.Random(SEED + 8)
    workbook = new_workbook()

    data = workbook.active
    data.title = "Data"
    write_header(data, 1, ["invoice_id", "supplier", "net", "vat_rate"])
    for index in range(200):
        row = index + 2
        data.cell(row=row, column=1, value=f"INV-{index + 1:06d}")
        data.cell(row=row, column=2, value=rng.choice(COUNTERPARTIES))
        data.cell(row=row, column=3, value=round(rng.uniform(50, 25_000), 2))
        data.cell(row=row, column=4, value=rng.choice([0.0, 0.05, 0.2]))

    calc = workbook.create_sheet("Calc")
    write_header(calc, 1, ["invoice_id", "net", "vat", "gross", "band"])
    for index in range(200):
        row = index + 2
        calc[f"A{row}"] = f"=Data!A{row}"
        calc[f"B{row}"] = f"=Data!C{row}"
        calc[f"C{row}"] = f"=ROUND(B{row}*Data!D{row},2)"
        calc[f"D{row}"] = f"=B{row}+C{row}"
        calc[f"E{row}"] = f'=IF(D{row}>10000,"large",IF(D{row}>1000,"medium","small"))'

    output = workbook.create_sheet("Output")
    write_header(output, 1, ["metric", "value"])
    output["A2"], output["B2"] = "total_net", "=ROUND(SUM(Calc!B2:B201),2)"
    output["A3"], output["B3"] = "total_vat", "=ROUND(SUM(Calc!C2:C201),2)"
    output["A4"], output["B4"] = "total_gross", "=ROUND(SUM(Calc!D2:D201),2)"
    output["A5"], output["B5"] = "large_count", '=COUNTIF(Calc!E2:E201,"large")'

    save(workbook, path)


# =============================================================================
# Verification — the hand-built parts are the ones most likely to be wrong
# =============================================================================


@dataclass(frozen=True, slots=True)
class Check:
    """One verification result."""

    fixture: str
    name: str
    detail: str


def verify_opens(path: Path) -> list[Check]:
    """Every fixture must open under both load modes without raising."""
    checks: list[Check] = []
    formulas = load_workbook(path)
    checks.append(
        Check(
            path.name,
            "opens (formulas)",
            f"{len(formulas.sheetnames)} sheets: {formulas.sheetnames}",
        )
    )
    values = load_workbook(path, data_only=True)
    checks.append(Check(path.name, "opens (data_only)", f"{len(values.sheetnames)} sheets"))
    formulas.close()
    values.close()
    return checks


def verify_cached_values(path: Path) -> list[Check]:
    """Assert real numbers come back from data_only, not None."""
    workbook = load_workbook(path, data_only=True)
    calc = workbook["Calc"]
    sampled = [calc[f"G{row}"].value for row in (2, 3, 250, 500, 501)]
    if any(value is None for value in sampled):
        msg = f"{path.name}: cached values missing on Calc!G — got {sampled}"
        raise AssertionError(msg)
    if not all(isinstance(value, (int, float)) for value in sampled):
        msg = f"{path.name}: cached values on Calc!G are not numeric — got {sampled}"
        raise AssertionError(msg)

    text = calc["A2"].value
    if not isinstance(text, str):
        msg = f"{path.name}: cached string value on Calc!A2 came back as {type(text).__name__}"
        raise AssertionError(msg)

    output = workbook["Output"]
    totals = [output[f"B{row}"].value for row in range(2, 8)]
    if any(value is None for value in totals):
        msg = f"{path.name}: cached values missing on Output!B — got {totals}"
        raise AssertionError(msg)

    # The formulas must survive the injection.
    with_formulas = load_workbook(path)
    formula = with_formulas["Calc"]["G2"].value
    if not isinstance(formula, str) or not formula.startswith("=ROUND("):
        msg = f"{path.name}: formula lost from Calc!G2 — got {formula!r}"
        raise AssertionError(msg)

    workbook.close()
    with_formulas.close()
    return [
        Check(path.name, "cached values present", f"Calc!G sample {sampled[:3]}"),
        Check(path.name, "cached strings typed", f"Calc!A2 = {text!r}"),
        Check(path.name, "formulas preserved", f"Calc!G2 = {formula!r}"),
        Check(path.name, "cached totals present", f"Output!B2 = {totals[0]}"),
    ]


def verify_no_cached_values(path: Path) -> list[Check]:
    """Assert every formula cell really does read None under data_only."""
    formulas = load_workbook(path)
    values = load_workbook(path, data_only=True)

    formula_cells = 0
    non_none = 0
    for sheet_name in formulas.sheetnames:
        source = formulas[sheet_name]
        cached = values[sheet_name]
        for row in source.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    formula_cells += 1
                    if cached[cell.coordinate].value is not None:
                        non_none += 1

    if formula_cells == 0:
        msg = f"{path.name}: expected formula cells, found none"
        raise AssertionError(msg)
    if non_none:
        msg = f"{path.name}: {non_none} of {formula_cells} formula cells have a cached value"
        raise AssertionError(msg)

    formulas.close()
    values.close()
    return [
        Check(path.name, "no cached values", f"all {formula_cells} formula cells read None"),
    ]


def verify_powerquery(path: Path) -> list[Check]:
    """Round-trip the DataMashup: locate, decode, unzip, read Section1.m back."""
    parts = read_parts(path)

    candidates = sorted(name for name in parts if re.fullmatch(r"customXml/item\d+\.xml", name))
    if len(candidates) < 3:
        msg = f"{path.name}: expected three customXml items, found {candidates}"
        raise AssertionError(msg)

    # Find by schema, exactly as the analyser must.
    found: str | None = None
    for name in candidates:
        if DATAMASHUP_NS.encode() in parts[name]:
            found = name
            break
    if found is None:
        msg = f"{path.name}: no customXml part declares the DataMashup namespace"
        raise AssertionError(msg)
    if found != "customXml/item3.xml":
        msg = f"{path.name}: DataMashup landed at {found}, expected customXml/item3.xml"
        raise AssertionError(msg)

    match = re.search(rb"<DataMashup[^>]*>(.*?)</DataMashup>", parts[found], re.DOTALL)
    if match is None:
        msg = f"{path.name}: could not find the DataMashup element in {found}"
        raise AssertionError(msg)

    envelope = base64.b64decode(match.group(1))
    blocks = parse_datamashup(envelope)

    with zipfile.ZipFile(io.BytesIO(blocks["package"])) as inner:
        names = inner.namelist()
        if "Formulas/Section1.m" not in names:
            msg = f"{path.name}: inner package has no Formulas/Section1.m, only {names}"
            raise AssertionError(msg)
        recovered = inner.read("Formulas/Section1.m").decode("utf-8")

    if recovered != SECTION_1_M:
        msg = f"{path.name}: recovered Section1.m does not match the source"
        raise AssertionError(msg)

    for query in POWER_QUERY_NAMES:
        if f"shared {query} =" not in recovered:
            msg = f"{path.name}: query {query!r} missing from the recovered M source"
            raise AssertionError(msg)

    if b"Excel.CurrentWorkbook" not in recovered.encode():
        msg = f"{path.name}: recovered M source has no Excel.CurrentWorkbook call"
        raise AssertionError(msg)

    # The envelope must genuinely carry the QDEFF header, not be a bare zip.
    if envelope[:4] != b"\x00\x00\x00\x00":
        msg = f"{path.name}: DataMashup envelope does not start with version 0"
        raise AssertionError(msg)
    if envelope[8:12] != b"PK\x03\x04":
        msg = f"{path.name}: inner package does not begin at offset 8"
        raise AssertionError(msg)

    decoys = [name for name in candidates if name != found]
    return [
        Check(path.name, "DataMashup located by schema", f"{found} (decoys: {decoys})"),
        Check(
            path.name,
            "QDEFF envelope parsed",
            f"package {len(blocks['package'])} bytes at offset 8",
        ),
        Check(
            path.name,
            "Section1.m round-trips",
            f"{len(recovered)} chars, queries {POWER_QUERY_NAMES}",
        ),
    ]


def verify_connections(path: Path) -> list[Check]:
    """Re-read connections.xml and assert the SQL survived intact."""
    parts = read_parts(path)
    if "xl/connections.xml" not in parts:
        msg = f"{path.name}: xl/connections.xml is missing"
        raise AssertionError(msg)

    if b"/xl/connections.xml" not in parts["[Content_Types].xml"]:
        msg = f"{path.name}: connections.xml has no content-type override"
        raise AssertionError(msg)
    if b"connections.xml" not in parts["xl/_rels/workbook.xml.rels"]:
        msg = f"{path.name}: connections.xml has no workbook relationship"
        raise AssertionError(msg)

    import xml.etree.ElementTree as ET

    root = ET.fromstring(parts["xl/connections.xml"].decode("utf-8"))
    connections = root.findall(f"{{{SHEET_MAIN_NS}}}connection")
    if len(connections) != 2:
        msg = f"{path.name}: expected two connections, found {len(connections)}"
        raise AssertionError(msg)

    recovered: dict[str, str] = {}
    for connection in connections:
        db = connection.find(f"{{{SHEET_MAIN_NS}}}dbPr")
        if db is None:
            msg = f"{path.name}: connection {connection.get('name')} has no dbPr"
            raise AssertionError(msg)
        recovered[connection.get("name", "")] = db.get("command", "")

    if recovered["RiskWarehouse"] != ODBC_SQL:
        msg = (
            f"{path.name}: ODBC SQL did not survive the round trip.\n"
            f"  expected {ODBC_SQL[:60]!r}...\n  got      {recovered['RiskWarehouse'][:60]!r}..."
        )
        raise AssertionError(msg)
    if recovered["FinanceCube"] != OLEDB_SQL:
        msg = f"{path.name}: OLEDB SQL did not survive the round trip"
        raise AssertionError(msg)

    # The newline encoding is the trap: a literal newline in an XML attribute is
    # normalised to a space by any conforming parser.
    if "\n" not in recovered["RiskWarehouse"]:
        msg = f"{path.name}: multi-line SQL came back as a single line"
        raise AssertionError(msg)

    lines = recovered["RiskWarehouse"].count("\n") + 1
    return [
        Check(path.name, "connections.xml wired", "content type + workbook relationship present"),
        Check(path.name, "two connections", "RiskWarehouse (ODBC), FinanceCube (OLEDB)"),
        Check(path.name, "multi-line SQL intact", f"{lines} lines recovered"),
    ]


def verify_hostile(path: Path) -> list[Check]:
    """Assert the traps are actually present in the saved file."""
    parts = read_parts(path)
    workbook = load_workbook(path)
    messy = workbook["Messy"]

    problems: list[str] = []
    if messy[f"D{HOSTILE_INCONSISTENT_ROW}"].value != (
        f"=B{HOSTILE_INCONSISTENT_ROW}*C{HOSTILE_INCONSISTENT_ROW}*1.1"
    ):
        problems.append(f"row {HOSTILE_INCONSISTENT_ROW} inconsistent formula missing")
    if messy["D46"].value != "=B46*C46":
        problems.append("uniform neighbours around row 47 missing")
    if messy["K9"].data_type != "e":
        problems.append("hard error cell K9 is not typed as an error")
    if not messy.merged_cells.ranges:
        problems.append("no merged cells")
    if not messy.column_dimensions["H"].hidden:
        problems.append("column H not hidden")
    if workbook["_Archive"].sheet_state != "hidden":
        problems.append("_Archive is not hidden")
    if "obsolete_rate_table" not in workbook.defined_names:
        problems.append("broken defined name missing")
    if not any(name.startswith("xl/externalLinks/") for name in parts):
        problems.append("external link part missing")
    if b"externalReferences" not in parts["xl/workbook.xml"]:
        problems.append("workbook.xml has no externalReferences")

    if problems:
        msg = f"{path.name}: hostile traps missing: {'; '.join(problems)}"
        raise AssertionError(msg)

    external = [name for name in parts if name.startswith("xl/externalLinks/")]
    workbook.close()
    return [
        Check(
            path.name,
            "inconsistent formula",
            f"Messy!D{HOSTILE_INCONSISTENT_ROW} breaks the region",
        ),
        Check(path.name, "error cell", "Messy!K9 has data_type 'e'"),
        Check(path.name, "hidden sheet and columns", "_Archive hidden; columns H, I hidden"),
        Check(path.name, "broken defined name", "obsolete_rate_table -> #REF!#REF!"),
        Check(path.name, "external link", f"{external}"),
    ]


def verify_word_companions(directory: Path) -> list[Check]:
    """Read the .docx back and confirm the .doc stub is not a real .doc."""
    from docx import Document

    docx_path = directory / "documented_procedure.docx"
    document = Document(str(docx_path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    if len(paragraphs) < 8:
        msg = f"{docx_path.name}: expected at least 8 paragraphs, found {len(paragraphs)}"
        raise AssertionError(msg)

    doc_path = directory / "procedure_legacy.doc"
    head = doc_path.read_bytes()[:8]
    if head == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        msg = f"{doc_path.name}: unexpectedly a real OLE2 compound file"
        raise AssertionError(msg)

    return [
        Check(docx_path.name, "docx readable", f"{len(paragraphs)} non-empty paragraphs"),
        Check(doc_path.name, "doc stub is not OLE2", f"leading bytes {head!r}"),
    ]


def verify_determinism(builders: dict[str, Any], tmp_root: Path) -> list[Check]:
    """Build every fixture twice and confirm the bytes agree."""
    first = tmp_root / "first"
    second = tmp_root / "second"
    checks: list[Check] = []
    for run in (first, second):
        run.mkdir(parents=True, exist_ok=True)
        for name, builder in builders.items():
            builder(run / name)
        build_word_companions(run)

    for name in sorted(builders):
        left = (first / name).read_bytes()
        right = (second / name).read_bytes()
        if left != right:
            msg = f"{name}: two runs of the generator produced different bytes"
            raise AssertionError(msg)
    checks.append(
        Check("(all)", "byte-deterministic", f"{len(builders)} workbooks built twice, identical")
    )

    for name in ("documented_procedure.docx", "procedure_legacy.doc"):
        if (first / name).read_bytes() != (second / name).read_bytes():
            msg = f"{name}: two runs of the generator produced different bytes"
            raise AssertionError(msg)
    checks.append(
        Check("(all)", "companions deterministic", "docx and doc stub identical across runs")
    )
    return checks


def verify_with_excel(path: Path) -> list[Check]:
    """Open the fixture in real Excel, recalculate, and compare cached values.

    This is the independent oracle for the hand-written ``<v>`` elements. It
    needs Excel and pywin32, so it is opt-in::

        uv run --with pywin32 python tests/fixtures/generate.py --verify-with-excel

    The workbook is opened read-only and never saved, so the committed fixture
    stays exactly as this script produced it.
    """
    try:
        import win32com.client
    except ImportError:
        return [Check(path.name, "excel oracle SKIPPED", "pywin32 is not installed")]

    cached = load_workbook(path, data_only=True)
    excel = win32com.client.Dispatch("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    mismatches: list[str] = []
    compared = 0
    try:
        book = excel.Workbooks.Open(str(path.resolve()), ReadOnly=True, UpdateLinks=0)
        excel.CalculateFullRebuild()
        for sheet_name, columns, rows in [
            ("Calc", (4, 5, 6, 7, 8), range(2, CLEAN_ROWS + 2)),
            ("Output", (2,), range(2, 8)),
        ]:
            sheet = book.Worksheets(sheet_name)
            for row in rows:
                for column in columns:
                    live = sheet.Cells(row, column).Value
                    stored = cached[sheet_name].cell(row=row, column=column).value
                    if live is None or stored is None:
                        mismatches.append(
                            f"{sheet_name}!R{row}C{column}: live={live} stored={stored}"
                        )
                        continue
                    compared += 1
                    if abs(float(live) - float(stored)) > 1e-9:
                        mismatches.append(
                            f"{sheet_name}!R{row}C{column}: Excel={live!r} fixture={stored!r}"
                        )
        book.Close(SaveChanges=False)
    finally:
        excel.Quit()
    cached.close()

    if mismatches:
        head = "\n  ".join(mismatches[:10])
        msg = (
            f"{path.name}: {len(mismatches)} cached values disagree with Excel's own "
            f"recalculation:\n  {head}"
        )
        raise AssertionError(msg)

    return [
        Check(path.name, "excel oracle agrees", f"{compared} cached values match Excel exactly"),
    ]


# =============================================================================
# Entry point
# =============================================================================

BUILDERS: dict[str, Any] = {
    "clean_pipeline.xlsx": build_clean_pipeline,
    "powerquery.xlsx": build_powerquery,
    "legacy_sql.xlsx": build_legacy_sql,
    "cross_sheet_chain.xlsx": build_cross_sheet_chain,
    "mostly_manual.xlsx": build_mostly_manual,
    "documented.xlsx": build_documented,
    "hostile.xlsx": build_hostile,
    "no_cached_values.xlsx": build_no_cached_values,
}


def generate(directory: Path) -> None:
    """Build every fixture and its companions into ``directory``."""
    directory.mkdir(parents=True, exist_ok=True)
    for name, builder in BUILDERS.items():
        builder(directory / name)
    build_word_companions(directory)


def verify(directory: Path) -> list[Check]:
    """Run every round-trip verification over an already-generated directory."""
    checks: list[Check] = []
    for name in BUILDERS:
        checks.extend(verify_opens(directory / name))
    checks.extend(verify_cached_values(directory / "clean_pipeline.xlsx"))
    checks.extend(verify_no_cached_values(directory / "no_cached_values.xlsx"))
    checks.extend(verify_powerquery(directory / "powerquery.xlsx"))
    checks.extend(verify_connections(directory / "legacy_sql.xlsx"))
    checks.extend(verify_hostile(directory / "hostile.xlsx"))
    checks.extend(verify_word_companions(directory))
    return checks


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument(
        "--out", type=Path, default=FIXTURE_DIR, help="directory to write fixtures into"
    )
    parser.add_argument(
        "--verify-with-excel",
        action="store_true",
        help="additionally recalculate clean_pipeline.xlsx in real Excel and compare",
    )
    parser.add_argument(
        "--check-determinism",
        action="store_true",
        help="build the corpus twice in a temporary directory and compare bytes",
    )
    args = parser.parse_args(argv)

    logging.basicConfig(level=logging.INFO, format="%(message)s")

    print(f"generating {len(BUILDERS)} fixtures into {args.out}")
    generate(args.out)

    checks = verify(args.out)

    if args.check_determinism:
        import tempfile

        with tempfile.TemporaryDirectory() as tmp:
            checks.extend(verify_determinism(BUILDERS, Path(tmp)))

    if args.verify_with_excel:
        checks.extend(verify_with_excel(args.out / "clean_pipeline.xlsx"))

    width = max(len(c.fixture) for c in checks)
    for check in checks:
        print(f"  ok  {check.fixture:<{width}}  {check.name}: {check.detail}")

    print(f"\n{len(checks)} checks passed")
    return 0


if __name__ == "__main__":
    sys.exit(main())
