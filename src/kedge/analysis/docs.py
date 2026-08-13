"""Process notes: documentation sheets, cell comments, and sibling documents.

The prose is worth as much as the formulas. A workbook that has been run monthly for six years
usually carries its real specification in three places — a tab called "Notes" that nobody has
opened since 2019, a scattering of cell comments explaining the exceptions, and a procedure
sitting next to the file on the shared drive — and none of it survives if the analyser only
reads cells.

Three sources, several shapes:

- **Documentation sheets** are stitched, not enumerated. One :class:`ProcessNote` per coherent
  block of prose, with a heading where the sheet has one; emitting a note per cell would turn
  a readable procedure into forty fragments and make the planner's job harder, not easier.
- **Cell comments** come through openpyxl, one note each, located by cell reference.
- **Sibling documents** are split on whatever structure their format actually carries. A
  ``.docx`` comes through python-docx and splits on the author's own heading styles; a ``.md``
  splits on its headings, which are the same statement in a different notation; a ``.txt``
  has no heading vocabulary at all and splits on blank lines, with ``heading`` left ``None``
  rather than guessed at. That difference is why ``markdown`` and ``plain_text`` are separate
  ``ProcessNote.source`` values: an absent heading is the author's silence in one and the
  format's silence in the other.

A sibling ``.doc`` -- the old binary format -- **fails loudly**: a ``doc_stub`` note and a
finding naming the file and the exact conversion command. Silently skipping it would leave the
user believing kedge had read a procedure it never opened (PLAN 1.5).

A document attached because its *filename* matched the convention, rather than because it
relates to the workbook's own name, earns a ``DOCUMENT_ATTACHED_BY_FILENAME`` finding. The
guess is usually right and is worth making, but both directions of it are otherwise invisible:
rename your procedure and it stops being read, and leave a colleague's ``SOP.docx`` in a shared
folder and it is read as though it described your workbook.

References:
- PLAN.md 1.5, 2.4 (documentation sheets), M1.
"""

from __future__ import annotations

import codecs
import logging
import posixpath
import re
import zipfile
from pathlib import Path
from typing import TYPE_CHECKING, Any
from xml.etree import ElementTree

from docx import Document
from openpyxl.utils import get_column_letter

from kedge.analysis.model import Finding, FindingKind, ProcessNote, Severity
from kedge.analysis.profile import workbook_view

if TYPE_CHECKING:
    from kedge.analysis.workbook import WorkbookHandle

logger = logging.getLogger(__name__)

__all__ = ["extract_notes", "sidecar_documents"]

MAX_NOTE_CHARS = 4_000
"""Per-note ceiling. A note longer than this is truncated with an explicit marker."""

_MAX_SHEETS = 100
_MAX_DOC_ROWS = 500
_MAX_DOC_COLS = 20
_MAX_NOTES_PER_SHEET = 50
_MAX_COMMENTS = 500
_MAX_SIDECARS = 10
"""Documents attached to one workbook, whatever their format. Shared deliberately: the cap
exists to bound what gets attached, and that concern does not change with the extension."""

_MAX_SIDECAR_NOTES = 200
"""Notes recovered from one sibling document. Also format-independent -- only the name was
ever specific to Word."""

_MAX_DOCX_TABLES = 20

_MAX_TEXT_BYTES = 2_000_000
"""Bytes read from a ``.md`` or ``.txt`` sidecar.

A ``.docx`` is bounded for free: python-docx reads a zip, and the paragraph loop stops at
``_MAX_SIDECAR_NOTES``. A text file has no such structure, so a 40MB one would be in memory
before any note cap could apply. Two megabytes is far past any real procedure and small
enough that ten of them are still bounded. Keep it a multiple of four: the truncation path
relies on that to avoid cutting a UTF-16 or UTF-32 code unit in half."""

_LONG_TEXT_CHARS = 40
"""A cell holding this much text is prose, not a label. The signal PLAN 2.4 classifies on."""

_MIN_BLOCK_CHARS = 25
_MAX_HEADING_CHARS = 80
_TRUNCATION_MARKER = " [... truncated]"

# `\b` is no use here: an underscore is a word character, so `\bprocedure\b` misses
# `procedure_legacy.doc` -- which is exactly how people name these files.
_DOC_SHEET_NAME_RE = re.compile(
    r"(?i)(?<![a-z])(read ?me|notes?|doc|docs|documentation|instructions?|process|procedure|"
    r"method|methodology|guide|about|overview|assumptions?|change ?log|history|control|"
    r"glossary)(?![a-z])"
)
_DOC_FILENAME_RE = re.compile(
    r"(?i)(?<![a-z])(process|procedure|method|notes?|readme|instructions?|guide|sop|runbook|"
    r"documentation|handover|spec|specification)(?![a-z])"
)

_WORD_SUFFIXES = (".docx", ".doc")
_TEXT_SUFFIXES = (".md", ".txt")
_SIDECAR_SUFFIXES = _WORD_SUFFIXES + _TEXT_SUFFIXES

# Markdown headings. ATX (`## Steps`) is the direct equivalent of a Word Heading style, and
# setext (a line underlined with `===` or `---`) is the other form the same author might use.
# Two dashes minimum: a lone `-` under a line is a list bullet at least as often as it is a
# heading, and reading a bullet as a heading silently restructures somebody's procedure.
_ATX_HEADING_RE = re.compile(r"^ {0,3}(?P<hashes>#{1,6})\s+(?P<text>.*?)\s*#*\s*$")
_SETEXT_UNDERLINE_RE = re.compile(r"^ {0,3}(=+|-{2,})\s*$")
_FENCE_RE = re.compile(r"^ {0,3}(```|~~~)")
"""A fenced code block suspends heading detection: `# rebuild the cache` inside a shell
snippet is a comment, and splitting a note on it would invent a section that is not there."""

# BOM to codec, widest first: the UTF-32 LE mark starts with the UTF-16 LE mark, so testing
# UTF-16 first would decode a UTF-32 file into interleaved nulls.
_BOM_ENCODINGS = (
    (codecs.BOM_UTF8, "utf-8-sig"),
    (codecs.BOM_UTF32_LE, "utf-32"),
    (codecs.BOM_UTF32_BE, "utf-32"),
    (codecs.BOM_UTF16_LE, "utf-16"),
    (codecs.BOM_UTF16_BE, "utf-16"),
)
_FALLBACK_ENCODINGS = ("utf-8", "cp1252")
"""No BOM, so guess. UTF-8 first because it is what everything modern writes; cp1252 second
because it is what Notepad wrote for twenty years and is still the Windows ANSI default."""

_TEXT_REMEDIATION = (
    "Re-save it as UTF-8 (Notepad: File > Save As > Encoding: UTF-8) and re-run kedge inspect."
)

_PATH_ATTRS = ("path", "workbook_path", "source_path", "filename", "file")
_ZIP_ATTRS = ("zf", "zip", "archive", "zip_file", "zipfile")
_COMMENTS_PART_RE = re.compile(r"(?i)^xl/comments.*\.xml$")
"""Excel writes ``xl/comments1.xml``; openpyxl writes ``xl/comments/comment1.xml``. Both count.
Only a pre-check -- the authoritative mapping comes from each sheet's relationships."""

_OLE2_SIGNATURE = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
_ZIP_SIGNATURE = b"PK\x03\x04"

_DOC_REMEDIATION = (
    "Convert it to .docx (Word: File > Save As > Word Document, or "
    "`soffice --headless --convert-to docx <file>`) and re-run kedge inspect."
)


# ── handle access ────────────────────────────────────────────────────────────────────────────


def _handle_path(handle: WorkbookHandle) -> Path | None:
    """The workbook's own path, which is what sibling documents are found relative to."""
    for attribute in _PATH_ATTRS:
        value = getattr(handle, attribute, None)
        if isinstance(value, Path):
            return value
        if isinstance(value, str) and value:
            return Path(value)
    return None


def _handle_zip(handle: WorkbookHandle) -> zipfile.ZipFile | None:
    for attribute in _ZIP_ATTRS:
        value = getattr(handle, attribute, None)
        if isinstance(value, zipfile.ZipFile):
            return value
    return None


def _worksheets(handle: WorkbookHandle) -> list[tuple[str, Any]]:
    """Named worksheets from the formula view, so a formula is distinguishable from prose."""
    book = workbook_view(handle, data_only=False) or workbook_view(handle, data_only=True)
    if book is None:
        logger.warning("no openpyxl view on the workbook handle; skipping in-sheet notes")
        return []

    named: list[tuple[str, Any]] = []
    for name in list(book.sheetnames)[:_MAX_SHEETS]:
        try:
            named.append((str(name), book[name]))
        except (KeyError, TypeError):
            logger.debug("sheet %r could not be opened for note extraction", name)
    return named


def _truncate(text: str) -> str:
    if len(text) <= MAX_NOTE_CHARS:
        return text
    return text[: MAX_NOTE_CHARS - len(_TRUNCATION_MARKER)] + _TRUNCATION_MARKER


# ── documentation sheets ─────────────────────────────────────────────────────────────────────


def _is_bold(cell: Any) -> bool:
    try:
        return bool(cell.font is not None and cell.font.bold)
    except Exception:  # a style openpyxl cannot resolve is not worth a traceback
        return False


class _Line:
    """One row of a documentation sheet, flattened to a single line of prose."""

    __slots__ = ("bold", "columns", "row", "text")

    def __init__(self, row: int, columns: list[int], text: str, *, bold: bool) -> None:
        self.row = row
        self.columns = columns
        self.text = text
        self.bold = bold


def _read_lines(worksheet: Any) -> tuple[list[_Line], int, int, int]:
    """Flatten a sheet into prose lines, and count what kind of sheet it looks like.

    Row and column numbers come from the iteration, never from the cell: in read-only mode a
    gap in a row is an ``EmptyCell``, which has no idea where it is.
    """
    max_row = min(int(getattr(worksheet, "max_row", 0) or 0), _MAX_DOC_ROWS)
    max_column = min(int(getattr(worksheet, "max_column", 0) or 0), _MAX_DOC_COLS)
    if max_row <= 0 or max_column <= 0:
        return [], 0, 0, 0

    lines: list[_Line] = []
    formula_cells = 0
    text_cells = 0
    long_text_cells = 0

    rows = worksheet.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_column)
    for row_number, row in enumerate(rows, start=1):
        pieces: list[str] = []
        columns: list[int] = []
        bold = False
        for column_number, cell in enumerate(row, start=1):
            value = cell.value
            if not isinstance(value, str):
                continue
            text = value.strip()
            if not text:
                continue
            if text.startswith("="):
                formula_cells += 1
                continue
            text_cells += 1
            if len(text) >= _LONG_TEXT_CHARS:
                long_text_cells += 1
            pieces.append(text)
            columns.append(column_number)
            bold = bold or _is_bold(cell)
        if pieces:
            lines.append(_Line(row_number, columns, " ".join(pieces), bold=bold))

    return lines, formula_cells, text_cells, long_text_cells


def _looks_documentary(
    title: str, formula_cells: int, text_cells: int, long_text_cells: int
) -> bool:
    """PLAN 2.4's documentation signal: mostly text, few or no formulas, long strings."""
    if text_cells == 0:
        return False
    if _DOC_SHEET_NAME_RE.search(title):
        return True
    return long_text_cells >= 3 and formula_cells * 4 <= text_cells


def _blocks(lines: list[_Line]) -> list[list[_Line]]:
    """Group lines into blocks, split wherever the sheet leaves a blank row."""
    grouped: list[list[_Line]] = []
    current: list[_Line] = []
    previous_row = None
    for line in lines:
        if previous_row is not None and line.row > previous_row + 1 and current:
            grouped.append(current)
            current = []
        current.append(line)
        previous_row = line.row
    if current:
        grouped.append(current)
    return grouped


def _looks_like_heading(line: _Line, block: list[_Line]) -> bool:
    if len(block) < 2:
        return False
    if len(line.text) > _MAX_HEADING_CHARS:
        return False
    return line.bold or not line.text.endswith((".", ":", ";", ","))


def _block_note(sheet_name: str, block: list[_Line]) -> ProcessNote | None:
    heading: str | None = None
    body = block
    if _looks_like_heading(block[0], block):
        heading = block[0].text
        body = block[1:]

    text = "\n".join(line.text for line in body).strip()
    if heading is None and len(text) < _MIN_BLOCK_CHARS:
        return None  # a stray label, not a note

    columns = [column for line in block for column in line.columns] or [1]
    first, last = block[0].row, block[-1].row
    start = f"{get_column_letter(min(columns))}{first}"
    end = f"{get_column_letter(max(columns))}{last}"
    return ProcessNote(
        source="sheet",
        origin=sheet_name,
        location=start if start == end else f"{start}:{end}",
        text=_truncate(text),
        heading=heading,
    )


def _sheet_notes(handle: WorkbookHandle) -> list[ProcessNote]:
    notes: list[ProcessNote] = []
    for title, worksheet in _worksheets(handle):
        try:
            lines, formula_cells, text_cells, long_text_cells = _read_lines(worksheet)
        except Exception:
            logger.warning("could not read sheet %r for process notes", title, exc_info=True)
            continue
        if not _looks_documentary(title, formula_cells, text_cells, long_text_cells):
            continue

        found = [
            note for block in _blocks(lines) if (note := _block_note(title, block)) is not None
        ]
        if found:
            logger.debug("stitched %d note(s) from documentation sheet %r", len(found), title)
        notes.extend(found[:_MAX_NOTES_PER_SHEET])
    return notes


# ── cell comments ────────────────────────────────────────────────────────────────────────────


def _part_names(handle: WorkbookHandle) -> frozenset[str]:
    names = getattr(handle, "part_names", None)
    if isinstance(names, (frozenset, set, list, tuple)):
        return frozenset(str(name) for name in names)
    archive = _handle_zip(handle)
    if archive is None:
        return frozenset()
    try:
        return frozenset(archive.namelist())
    except (OSError, RuntimeError, zipfile.BadZipFile):
        return frozenset()


def _read_part(handle: WorkbookHandle, name: str) -> bytes | None:
    reader = getattr(handle, "read_part", None)
    if callable(reader):
        try:
            return reader(name)
        except Exception:
            logger.debug("could not read part %s", name, exc_info=True)
            return None
    archive = _handle_zip(handle)
    if archive is None:
        return None
    try:
        return archive.read(name)
    except (KeyError, OSError, RuntimeError, zipfile.BadZipFile):
        return None


def _resolve_target(base_part: str, target: str) -> str:
    """Resolve a relationship target against the part that declared it."""
    if target.startswith("/"):
        return target.lstrip("/")
    return posixpath.normpath(posixpath.join(posixpath.dirname(base_part), target))


def _comment_parts(handle: WorkbookHandle) -> list[tuple[str, str]]:
    """Map each sheet to its comments part, via that sheet's relationships.

    Read-only openpyxl drops cell comments entirely, so they are recovered from the archive.
    The sheet-to-part mapping comes from :meth:`WorkbookHandle.structure`, which already owns
    it -- deriving it a second time here would be two things to keep in step.
    """
    names = _part_names(handle)
    if not any(_COMMENTS_PART_RE.match(name) for name in names):
        logger.debug("workbook carries no comment parts")
        return []

    pairs: list[tuple[str, str]] = []
    for sheet in list(getattr(handle, "sheet_names", []))[:_MAX_SHEETS]:
        structure = handle.structure(sheet)
        part = getattr(structure, "part", None)
        if not part:
            continue
        rels_path = f"{posixpath.dirname(part)}/_rels/{posixpath.basename(part)}.rels"
        data = _read_part(handle, rels_path)
        if not data:
            continue
        try:
            root = ElementTree.fromstring(data)
        except ElementTree.ParseError:
            logger.debug("relationships for sheet %r are not well-formed", sheet)
            continue
        for relationship in root:
            rel_type = relationship.get("Type", "")
            target = relationship.get("Target", "")
            if not target or rel_type.rsplit("/", maxsplit=1)[-1] != "comments":
                continue
            pairs.append((sheet, _resolve_target(part, target)))
    return pairs


def _strip_threaded_preamble(text: str) -> str:
    """Drop the compatibility preamble Excel writes into a threaded comment's legacy mirror."""
    if not text.lstrip().startswith("[Threaded comment]"):
        return text
    _, marker, body = text.partition("Comment:")
    return body.strip() if marker else text


def _comments_from_part(data: bytes, sheet: str) -> list[ProcessNote]:
    root = ElementTree.fromstring(data)
    authors = [
        (element.text or "").strip()
        for element in root.iter()
        if element.tag.rsplit("}", maxsplit=1)[-1] == "author"
    ]

    notes: list[ProcessNote] = []
    for element in root.iter():
        if element.tag.rsplit("}", maxsplit=1)[-1] != "comment":
            continue
        runs = [
            node.text or ""
            for node in element.iter()
            if node.tag.rsplit("}", maxsplit=1)[-1] == "t"
        ]
        text = _strip_threaded_preamble("".join(runs).strip())
        if not text:
            continue
        author_index = element.get("authorId", "")
        author = (
            authors[int(author_index)]
            if author_index.isdigit() and int(author_index) < len(authors)
            else None
        )
        notes.append(
            ProcessNote(
                source="cell_comment",
                origin=sheet,
                location=element.get("ref") or None,
                text=_truncate(text),
                heading=author or None,
            )
        )
    return notes


def _comment_notes(handle: WorkbookHandle) -> list[ProcessNote]:
    notes: list[ProcessNote] = []
    for sheet, part in _comment_parts(handle):
        data = _read_part(handle, part)
        if not data:
            continue
        try:
            found = _comments_from_part(data, sheet)
        except ElementTree.ParseError as exc:
            logger.warning("could not parse %s: %s", part, exc)
            continue
        notes.extend(found)
        if len(notes) >= _MAX_COMMENTS:
            logger.warning("stopping at %d cell comments", _MAX_COMMENTS)
            return notes[:_MAX_COMMENTS]
    return notes


# ── sibling Word documents ───────────────────────────────────────────────────────────────────


def _sidecar_matches(workbook_path: Path) -> list[tuple[Path, bool]]:
    """Sibling documents, each flagged with whether only its filename argued for it.

    The flag is the whole of :data:`FindingKind.DOCUMENT_ATTACHED_BY_FILENAME`: a stem match
    is a statement about *this* workbook and needs no warning, while a filename match is a
    guess made on a naming convention and is worth recording. Computed here rather than
    recomputed by the caller so the two can never drift.
    """
    directory = workbook_path.parent
    if not directory.is_dir():
        return []

    stem = workbook_path.stem.casefold()
    matches: list[tuple[Path, bool]] = []
    try:
        entries = sorted(directory.iterdir())
    except OSError as exc:
        logger.warning("could not list %s for sibling documents: %s", directory, exc)
        return []

    for entry in entries:
        if entry.suffix.casefold() not in _SIDECAR_SUFFIXES or entry.name.startswith("~$"):
            continue
        candidate = entry.stem.casefold()
        related = candidate.startswith(stem) or stem.startswith(candidate)
        if related:
            matches.append((entry, False))
        elif _DOC_FILENAME_RE.search(entry.stem):
            matches.append((entry, True))
    return matches[:_MAX_SIDECARS]


def sidecar_documents(workbook_path: Path) -> list[Path]:
    """Find documents beside the workbook that plausibly document it.

    Matched two ways: a file whose stem relates to the workbook's own (``rwa_monthly.xlsx``
    and ``rwa_monthly - procedure.docx``), or one named conventionally (``Process Notes``,
    ``SOP``, ``Runbook``). Deliberately not "every document in the folder" — on a shared drive
    that attaches somebody else's procedure to this workbook's analysis.

    Args:
        workbook_path: Path of the workbook being analysed.

    Returns:
        Matching ``.docx``, ``.doc``, ``.md`` and ``.txt`` paths, sorted, capped at ten. Word
        lock files (``~$...``) are excluded.
    """
    return [path for path, _ in _sidecar_matches(workbook_path)]


def _docx_notes(path: Path) -> list[ProcessNote]:
    """Read a .docx into notes, one per heading section plus one per table."""
    document = Document(str(path))
    notes: list[ProcessNote] = []
    origin = str(path)
    heading: str | None = None
    buffer: list[str] = []
    block_start = 1

    def flush(end: int) -> None:
        text = "\n".join(buffer).strip()
        if not text:
            return
        location = (
            f"paragraph {block_start}" if block_start == end else f"paragraphs {block_start}-{end}"
        )
        notes.append(
            ProcessNote(
                source="docx",
                origin=origin,
                location=location,
                text=_truncate(text),
                heading=heading,
            )
        )

    for number, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        style = ""
        try:
            style = str(paragraph.style.name or "") if paragraph.style is not None else ""
        except Exception:  # a document with a broken style table still has readable text
            logger.debug("could not read a paragraph style in %s", path, exc_info=True)
        if style.startswith("Heading") or style in ("Title", "Subtitle"):
            flush(number - 1)
            buffer = []
            heading = text
            block_start = number
            continue
        if not buffer:
            block_start = number
        buffer.append(text)
        if len(notes) >= _MAX_SIDECAR_NOTES:
            break
    flush(len(document.paragraphs))

    for index, table in enumerate(document.tables[:_MAX_DOCX_TABLES], start=1):
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            for row in table.rows
        ]
        text = "\n".join(line for line in rows if line).strip()
        if text:
            notes.append(
                ProcessNote(
                    source="docx",
                    origin=origin,
                    location=f"table {index}",
                    text=_truncate(text),
                    heading=heading,
                )
            )
    return notes


# ── sibling markdown and plain text ──────────────────────────────────────────────────────────


def _drop_partial_tail(raw: bytes) -> bytes:
    """Drop an incomplete UTF-8 sequence left dangling by the byte cap.

    Only ever applied to a buffer kedge itself truncated, so a broken character at the end is
    kedge's doing and not the file's. It can also take one *complete* final character with it;
    on a buffer that has already lost the rest of the file, that is not worth the arithmetic
    to distinguish.
    """
    index = len(raw)
    while index and len(raw) - index < 3 and 0x80 <= raw[index - 1] < 0xC0:
        index -= 1
    if index and raw[index - 1] >= 0xC0:
        return raw[: index - 1]
    return raw


def _decode_text(raw: bytes) -> str | None:
    """Decode a text sidecar, byte-order mark first, then UTF-8, then Windows-1252.

    Returns ``None`` when nothing in the ladder produces text, which is a :class:`Finding`
    rather than an exception. Note that ``UnicodeDecodeError`` is a ``ValueError`` and not an
    ``OSError``: letting one escape would fall straight past the caller's ``except OSError``
    to the catch-all in :func:`extract_notes`, costing every sidecar note on the workbook
    rather than the one file that could not be read.
    """
    for bom, encoding in _BOM_ENCODINGS:
        if raw.startswith(bom):
            try:
                return raw.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                return None

    for encoding in _FALLBACK_ENCODINGS:
        try:
            text = raw.decode(encoding)
        except UnicodeDecodeError:
            continue
        if "\x00" in text:
            # cp1252 decodes almost any byte, so "it decoded" is not "it is text". A null
            # means UTF-16 with no mark, read one byte at a time into mojibake.
            return None
        return text
    return None


def _read_text_lines(path: Path) -> tuple[list[str], Finding | None]:
    """Read a text sidecar into lines, bounded by :data:`_MAX_TEXT_BYTES`."""
    try:
        with path.open("rb") as stream:
            raw = stream.read(_MAX_TEXT_BYTES + 1)
    except OSError as exc:
        logger.warning("could not open %s: %s", path, exc)
        return [], None

    truncated = len(raw) > _MAX_TEXT_BYTES
    if truncated:
        # A multiple of four never splits a UTF-16 or UTF-32 code unit; the UTF-8 tail is
        # trimmed separately. `stat()` is not consulted -- a file being written grows between
        # the stat and the read, and the read is the only figure that binds.
        logger.warning("%s exceeds %d bytes; reading the first part only", path, _MAX_TEXT_BYTES)
        raw = _drop_partial_tail(raw[:_MAX_TEXT_BYTES])

    text = _decode_text(raw)
    if text is None:
        logger.warning("could not decode %s as text", path)
        return [], Finding(
            kind=FindingKind.UNPARSEABLE_PART,
            severity=Severity.WARNING,
            message=(
                f"could not read the text document '{path.name}': it is not UTF-8, "
                f"Windows-1252, or any encoding it declares a byte-order mark for. "
                f"Any process notes in it are missing from this analysis."
            ),
            location=str(path),
            remediation=_TEXT_REMEDIATION,
        )

    lines = text.splitlines()
    if truncated:
        if len(lines) > 1:
            lines = lines[:-1]  # the cap landed mid-line; half a sentence is worse than none
        lines.append(_TRUNCATION_MARKER.strip())
    return lines, None


def _location_lines(start: int, end: int) -> str:
    """The honest equivalent of the Word reader's ``paragraphs 4-11``.

    Line numbers, because that is how both formats are addressed everywhere else -- an editor,
    a diff, a review comment. One-based, counting blank lines, so the number matches what the
    reader sees when they open the file.
    """
    return f"line {start}" if start == end else f"lines {start}-{end}"


def _markdown_tokens(lines: list[str]) -> list[tuple[int, bool, str]]:
    """``(line number, is a heading, text)``, with setext underlines folded into their heading.

    Resolving setext here rather than in the note loop keeps that loop the same shape as the
    Word reader's: one pass, flush on a heading.
    """
    tokens: list[tuple[int, bool, str]] = []
    fenced = False
    skip_next = False

    for index, line in enumerate(lines):
        if skip_next:
            skip_next = False
            continue
        if _FENCE_RE.match(line):
            fenced = not fenced
            continue

        number = index + 1
        text = line.strip()
        if not text:
            continue
        if fenced:
            tokens.append((number, False, text))
            continue

        atx = _ATX_HEADING_RE.match(line)
        if atx:
            heading = atx.group("text").strip()
            if heading:
                tokens.append((number, True, heading))
            continue

        following = lines[index + 1] if index + 1 < len(lines) else ""
        if _SETEXT_UNDERLINE_RE.match(following):
            tokens.append((number, True, text))
            skip_next = True
            continue

        tokens.append((number, False, text))
    return tokens


def _markdown_notes(path: Path, lines: list[str]) -> list[ProcessNote]:
    """Split markdown on its own headings, exactly as the Word reader splits on styles."""
    notes: list[ProcessNote] = []
    origin = str(path)
    heading: str | None = None
    buffer: list[str] = []
    block_start = 0
    block_end = 0

    def flush() -> None:
        text = "\n".join(buffer).strip()
        if not text:
            return
        notes.append(
            ProcessNote(
                source="markdown",
                origin=origin,
                location=_location_lines(block_start, block_end),
                text=_truncate(text),
                heading=heading,
            )
        )

    for number, is_heading, text in _markdown_tokens(lines):
        if is_heading:
            flush()
            if len(notes) >= _MAX_SIDECAR_NOTES:
                return notes
            buffer = []
            heading = text
            block_start = number
            block_end = number
            continue
        if not buffer:
            block_start = number
        block_end = number
        buffer.append(text)
    flush()
    return notes[:_MAX_SIDECAR_NOTES]


def _plain_text_notes(path: Path, lines: list[str]) -> list[ProcessNote]:
    """Split plain text on blank lines, which is the only structure it has.

    ``heading`` stays ``None`` throughout. A ``.txt`` carries no heading vocabulary, and the
    sheet reader's guess -- short, not ending in punctuation -- keys on bold text a file does
    not have and would be a coin flip on the first line of every paragraph. A guess here would
    reach the planning prompt as though the author had written it, so there is none.

    Unlike the sheet reader, a short block is kept: ``_MIN_BLOCK_CHARS`` exists there to drop
    a stray label in a spreadsheet cell, and nothing in a text file is stray. "Run the model."
    is a step, and dropping it would be the silence this whole extractor exists to end.
    """
    notes: list[ProcessNote] = []
    origin = str(path)
    buffer: list[str] = []
    block_start = 0
    block_end = 0

    def flush() -> None:
        text = "\n".join(buffer).strip()
        if not text:
            return
        notes.append(
            ProcessNote(
                source="plain_text",
                origin=origin,
                location=_location_lines(block_start, block_end),
                text=_truncate(text),
                heading=None,
            )
        )

    for index, line in enumerate(lines):
        text = line.strip()
        if not text:
            flush()
            if len(notes) >= _MAX_SIDECAR_NOTES:
                return notes
            buffer = []
            continue
        if not buffer:
            block_start = index + 1
        block_end = index + 1
        buffer.append(text)
    flush()
    return notes[:_MAX_SIDECAR_NOTES]


def _text_sidecar_notes(path: Path) -> tuple[list[ProcessNote], Finding | None]:
    """Read a ``.md`` or ``.txt`` sidecar. Never raises past this boundary."""
    lines, finding = _read_text_lines(path)
    if finding is not None or not lines:
        return [], finding

    try:
        if path.suffix.casefold() == ".md":
            return _markdown_notes(path, lines), None
        return _plain_text_notes(path, lines), None
    except Exception as exc:  # non-negotiable 4: a note extractor never costs the analysis
        logger.warning("could not split %s into notes: %s", path, exc, exc_info=True)
        return [], Finding(
            kind=FindingKind.UNPARSEABLE_PART,
            severity=Severity.WARNING,
            message=f"could not read the text document '{path.name}': {exc}",
            location=str(path),
            remediation=_TEXT_REMEDIATION,
        )


# ── sibling document findings ────────────────────────────────────────────────────────────────


def _attached_by_filename(path: Path, workbook_path: Path) -> Finding:
    """Record a document attached on the strength of its name alone."""
    message = (
        f"'{path.name}' was attached to '{workbook_path.name}' because its name matches the "
        f"documentation naming convention, not because it relates to the workbook's own name. "
        f"Its contents are being read as though they described this workbook."
    )
    return Finding(
        kind=FindingKind.DOCUMENT_ATTACHED_BY_FILENAME,
        severity=Severity.INFO,
        message=message,
        location=str(path),
        remediation=(
            f"If it does describe this workbook, rename it to "
            f"'{workbook_path.stem} - {path.stem}{path.suffix}' so the match is deliberate. "
            f"If it does not, move it out of {path.parent} so it stops being read."
        ),
    )


def _doc_stub(path: Path) -> tuple[ProcessNote, Finding]:
    """A legacy .doc: say so, name the file, and give the conversion command."""
    message = (
        f"'{path.name}' sits beside the workbook but is in the legacy .doc binary format, "
        f"which kedge cannot read. Any process notes in it are missing from this analysis."
    )
    note = ProcessNote(
        source="doc_stub",
        origin=str(path),
        location=None,
        text=f"{message} {_DOC_REMEDIATION}",
        heading=None,
    )
    finding = Finding(
        kind=FindingKind.UNSUPPORTED_FORMAT,
        severity=Severity.WARNING,
        message=message,
        location=str(path),
        remediation=_DOC_REMEDIATION,
    )
    return note, finding


def _sidecar_notes(workbook_path: Path) -> tuple[list[ProcessNote], list[Finding]]:
    notes: list[ProcessNote] = []
    findings: list[Finding] = []

    for path, by_filename in _sidecar_matches(workbook_path):
        if by_filename:
            logger.debug("%s matched the naming convention rather than the workbook stem", path)
            findings.append(_attached_by_filename(path, workbook_path))

        if path.suffix.casefold() in _TEXT_SUFFIXES:
            found, finding = _text_sidecar_notes(path)
            if finding is not None:
                findings.append(finding)
            else:
                logger.info("read %d process note(s) from %s", len(found), path.name)
            notes.extend(found)
            continue

        try:
            with path.open("rb") as handle:
                header = handle.read(8)
        except OSError as exc:
            logger.warning("could not open %s: %s", path, exc)
            continue

        if path.suffix.casefold() == ".doc" and not header.startswith(_ZIP_SIGNATURE):
            if not header.startswith(_OLE2_SIGNATURE):
                logger.debug("%s is neither OLE2 nor a zip; reporting it as unreadable", path)
            note, finding = _doc_stub(path)
            notes.append(note)
            findings.append(finding)
            continue

        try:
            found = _docx_notes(path)
        except Exception as exc:  # python-docx raises a wide family on a damaged package
            logger.warning("could not read %s: %s", path, exc)
            findings.append(
                Finding(
                    kind=FindingKind.UNPARSEABLE_PART,
                    severity=Severity.WARNING,
                    message=f"could not read the Word document '{path.name}': {exc}",
                    location=str(path),
                    remediation=(
                        "Open it in Word and re-save it. If it will not open there either, "
                        "the file is damaged and its process notes are lost."
                    ),
                )
            )
            continue

        logger.info("read %d process note(s) from %s", len(found), path.name)
        notes.extend(found)

    return notes, findings


# ── entry point ──────────────────────────────────────────────────────────────────────────────


def extract_notes(handle: WorkbookHandle) -> tuple[list[ProcessNote], list[Finding]]:
    """Collect every piece of process prose attached to a workbook.

    Draws on documentation sheets, cell comments, and sibling documents, in that order. Each
    source is isolated: an unreadable sheet, a workbook openpyxl will not walk, a damaged
    ``.docx`` or a ``.txt`` in an encoding nothing will decode costs that source alone and
    produces a warning or a finding, never an exception (CONVENTIONS non-negotiable 4).

    Args:
        handle: The open workbook handle. Its ``path`` is what sibling documents are found
            relative to; without one, only in-workbook sources are read.

    Returns:
        The notes found, in source order, and any findings raised while reading them. A
        workbook with no prose anywhere returns two empty lists and that is a normal result.
    """
    notes: list[ProcessNote] = []
    findings: list[Finding] = []

    for source, reader in (("sheet", _sheet_notes), ("cell_comment", _comment_notes)):
        try:
            notes.extend(reader(handle))
        except Exception:
            logger.warning("the %s note extractor failed; continuing", source, exc_info=True)

    workbook_path = _handle_path(handle)
    if workbook_path is None:
        logger.debug("workbook handle exposes no path; skipping sibling documents")
    else:
        try:
            sidecar_notes, sidecar_findings = _sidecar_notes(workbook_path)
        except Exception:
            logger.warning("could not scan for sibling documents", exc_info=True)
        else:
            notes.extend(sidecar_notes)
            findings.extend(sidecar_findings)

    logger.info("extracted %d process note(s)", len(notes))
    return notes, findings
